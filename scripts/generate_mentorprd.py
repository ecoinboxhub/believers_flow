#!/usr/bin/env python3
"""Generate mentorprd.docx - BelieversFlow PRD for Mentor Presentation"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ---- Page Setup ----
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ---- Styles ----
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for i in range(1, 4):
    heading = doc.styles[f'Heading {i}']
    heading.font.color.rgb = RGBColor(0x1a, 0x0a, 0x2e)
    heading.font.name = 'Calibri'

# ---- Title Page ----
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BelieversFlow')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x0a, 0x2e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Product Requirements Document')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x7b, 0x2d, 0x8e)

doc.add_paragraph()

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run('Version 4.2.0')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0xf2, 0xc9, 0x4c)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Prepared for Mentor Review\nJuly 4, 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

conf = doc.add_paragraph()
conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = conf.add_run('Classification: Confidential')
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ---- Table of Contents ----
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Problem Statement',
    '3. Competitor Analysis',
    '4. Product Vision & Objectives',
    '5. Target Audience',
    '6. Implementation Status',
    '7. Updates & Upgrades So Far',
    '8. Expected Upgrades',
    '9. Technical Architecture',
    '10. Legal & Compliance Framework',
    '11. Risk Assessment',
    '12. Next Steps',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ---- 1. Executive Summary ----
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'BelieversFlow is a single-page, offline-first Christian productivity application that unifies '
    'task management, prayer tracking, Bible reading, journaling, and AI-powered faith guidance into '
    'one seamless experience. Unlike generic productivity tools that ignore spiritual dimensions, '
    'BelieversFlow is purpose-built for Christians who want to integrate their faith into their daily planning.'
)

doc.add_paragraph()
doc.add_heading('Key Facts', level=2)

facts = [
    ('Version', '4.2.0'),
    ('Stack', 'React 19 + Vite 8 + Capacitor 8 + Python FastAPI + PostgreSQL + Pinecone RAG + Multi-LLM'),
    ('Platforms', 'Web (PWA), Android APK, iOS (planned)'),
    ('Revenue Model', 'Freemium (all features free for logged-in users)'),
    ('License', 'MIT (open source)'),
    ('Tests', '61/61 passing (Vitest unit tests)'),
    ('Backend', '4.2.0 with auth, sync, RAG, multi-LLM, Google OAuth, legal tracking'),
    ('Database', 'PostgreSQL on Aiven with asyncpg'),
    ('Vector Store', 'Pinecone (54 Bible verses, 1024-dim embeddings)'),
    ('Legal Framework', '14 documents implemented'),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = 'Fact'
hdr[1].text = 'Value'
for fact, value in facts:
    row = table.add_row().cells
    row[0].text = fact
    row[1].text = value

doc.add_paragraph()
doc.add_heading('Production Status (July 4, 2026)', level=2)

status_items = [
    'User authentication (email + Google OAuth)',
    'Cloud sync (last-write-wins conflict resolution)',
    'Premium feature gating (freemium model)',
    'Pinecone RAG for Bible verse search',
    'Multi-LLM support (GROQ, OpenAI, OpenRouter)',
    'PostgreSQL database on Aiven',
    '61 unit tests passing',
    'All UI tabs visible and functional',
    'Legal & compliance framework (14 documents)',
    'In-app legal acceptance flow',
    'Settings Legal tab',
    'Security audit fixes (credentials removed, consistent dimensions)',
]

for item in status_items:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.text = item

doc.add_page_break()

# ---- 2. Problem Statement ----
doc.add_heading('2. Problem Statement', level=1)

doc.add_heading('2.1 The Core Problem', level=2)
doc.add_paragraph(
    'Christians lack an integrated digital tool that serves both their productivity needs and '
    'spiritual growth. Existing solutions force users to juggle multiple disconnected apps:'
)

problems = [
    ('Todoist / TickTick', 'Task management (secular, no faith context)'),
    ('YouVersion', 'Bible reading (no task integration)'),
    ('Notes app / physical journal', 'Prayer journal (no structure or streak)'),
    ('Notion / Google Docs', 'Bible study notes (overkill, no mobile-first design)'),
]

for name, desc in problems:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    run = p.add_run(f'{name} — ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('2.2 Pain Points', level=2)

pain_table = doc.add_table(rows=1, cols=3)
pain_table.style = 'Light Grid Accent 1'
hdr = pain_table.rows[0].cells
hdr[0].text = '#'
hdr[1].text = 'Pain Point'
hdr[2].text = 'Impact'

pains = [
    ('P1', 'No unified view of daily tasks + spiritual activities', 'Context switching, missed prayers'),
    ('P2', 'No prayer discipline tracking', 'Inconsistent quiet time'),
    ('P3', 'Privacy concerns with cloud-based journals', 'Users self-censor'),
    ('P4', 'Spiritual vs. secular balance is invisible', 'Unaware of time allocation'),
    ('P5', 'Bible reading plans exist separately', 'Reading happens inconsistently'),
    ('P6', 'No AI guidance tailored to Christian context', 'Users turn to generic AI'),
]

for num, point, impact in pains:
    row = pain_table.add_row().cells
    row[0].text = num
    row[1].text = point
    row[2].text = impact

doc.add_page_break()

# ---- 3. Competitor Analysis ----
doc.add_heading('3. Competitor Analysis', level=1)

doc.add_heading('3.1 Direct Competitors (Faith + Productivity)', level=2)

comp_table = doc.add_table(rows=1, cols=5)
comp_table.style = 'Light Grid Accent 1'
hdr = comp_table.rows[0].cells
hdr[0].text = 'Competitor'
hdr[1].text = 'Revenue Model'
hdr[2].text = 'Strengths'
hdr[3].text = 'Weaknesses'
hdr[4].text = 'BF Advantage'

competitors = [
    ('YouVersion', 'Free (donations)', '500M+ installs, Bible tools', 'No tasks, no prayer, requires account', 'Integrated faith + productivity'),
    ('Pray.com', 'Freemium ($4.99/mo)', 'Audio prayers, meditation', 'No tasks, not offline', 'Full features free, offline-first'),
    ('Abide', 'Subscription ($11.99/mo)', 'Sleep stories, high quality', 'Not productivity, expensive', 'Broader features, free'),
    ('Glo Bible', 'Paid app ($4.99)', 'Rich media Bible', 'No tasks/prayer, paid upfront', 'Free, AI-powered tools'),
    ('Bible Gateway', 'Free / Ads', 'Web Bible lookup', 'No mobile app, ad-supported', 'Native app, offline, no ads'),
]

for name, rev, strength, weakness, advantage in competitors:
    row = comp_table.add_row().cells
    row[0].text = name
    row[1].text = rev
    row[2].text = strength
    row[3].text = weakness
    row[4].text = advantage

doc.add_heading('3.2 Feature Comparison Matrix', level=2)

matrix_table = doc.add_table(rows=1, cols=7)
matrix_table.style = 'Light Grid Accent 1'
hdr = matrix_table.rows[0].cells
hdr[0].text = 'Feature'
hdr[1].text = 'BelieversFlow'
hdr[2].text = 'YouVersion'
hdr[3].text = 'Todoist'
hdr[4].text = 'Pray.com'
hdr[5].text = 'Day One'
hdr[6].text = 'Notion'

features = [
    ('Task management', '✅ Free', '❌', '✅ Freemium', '❌', '❌', '✅ Freemium'),
    ('Prayer tracking', '✅ Free', '❌', '❌', '✅ Limited', '❌', '❌'),
    ('Bible reader (66 books)', '✅ Free', '✅ Free', '❌', '❌', '❌', '❌'),
    ('AI faith assistant', '✅ Free', '❌', '❌', '❌', '❌', '❌'),
    ('Journal / diary', '✅ Free', '❌', '❌', '❌', '✅ Sub', '✅ Freemium'),
    ('Hymn book (1,001 hymns)', '✅ Free', '❌', '❌', '❌', '❌', '❌'),
    ('Daily devotional', '✅ Free', '✅ Free', '❌', '✅ Sub', '❌', '❌'),
    ('Offline-first', '✅ Full', '⚠️ Partial', '⚠️ Partial', '❌', '⚠️ Partial', '⚠️ Partial'),
    ('No account required', '✅', '❌', '❌', '❌', '❌', '❌'),
    ('Open source', '✅', '❌', '❌', '❌', '❌', '❌'),
    ('Customization', '✅ 5+ themes', '⚠️ Limited', '⚠️ Premium', '⚠️ Limited', '⚠️ Limited', '⚠️ Limited'),
    ('Price (full features)', 'FREE', 'Free', '$4/mo', '$4.99/mo', '$2.99/mo', '$10/mo'),
]

for row_data in features:
    row = matrix_table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

doc.add_heading('3.3 Competitive Moat', level=2)

moats = [
    ('Integration', 'No competitor combines tasks + prayer + Bible + diary + hymns + AI in one app'),
    ('Privacy', 'Zero data collection, no accounts required, open source — verifiable by users'),
    ('Offline-first', 'Full functionality without internet; competitors degrade significantly offline'),
    ('Price', '40+ features completely free; competitors gate similar features behind subscriptions'),
    ('Open source', 'Community trust; anyone can audit the code; fork-friendly'),
]

for name, desc in moats:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    run = p.add_run(f'{name}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ---- 4. Product Vision & Objectives ----
doc.add_heading('4. Product Vision & Objectives', level=1)

doc.add_heading('4.1 Vision Statement', level=2)
vision = doc.add_paragraph()
vision.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = vision.add_run(
    '"To be the most trusted digital companion for daily Christian living — combining productivity, '
    'scripture, prayer, and AI guidance in a private, offline-first experience."'
)
run.italic = True
run.font.size = Pt(12)

doc.add_heading('4.2 Product Objectives', level=2)

obj_table = doc.add_table(rows=1, cols=4)
obj_table.style = 'Light Grid Accent 1'
hdr = obj_table.rows[0].cells
hdr[0].text = 'Objective'
hdr[1].text = 'Metric'
hdr[2].text = 'Target'
hdr[3].text = 'Status'

objectives = [
    ('O1: Daily Active Users', 'DAU', '1,000 by Q4 2026', 'Measuring'),
    ('O2: User Retention', '7-day retention', '>40%', 'Measuring'),
    ('O3: Prayer Streak', 'Weekly prayer logging', '>50% of DAU', 'Measuring'),
    ('O4: User Satisfaction', 'Rating', '>4.5/5', 'Measuring'),
    ('O5: GitHub Stars', 'Stars', '>100 by Q4 2026', 'Measuring'),
    ('O6: Test Coverage', 'Coverage', '>80% of pure functions', '61 tests passing'),
    ('O7: Lighthouse', 'Mobile score', '>80', 'Measuring'),
]

for obj, metric, target, status in objectives:
    row = obj_table.add_row().cells
    row[0].text = obj
    row[1].text = metric
    row[2].text = target
    row[3].text = status

doc.add_heading('4.3 Key Differentiators', level=2)

diff_table = doc.add_table(rows=1, cols=3)
diff_table.style = 'Light Grid Accent 1'
hdr = diff_table.rows[0].cells
hdr[0].text = 'Differentiator'
hdr[1].text = 'BelieversFlow'
hdr[2].text = 'Competitors'

diffs = [
    ('Integrated faith + productivity', '✅ Tasks + prayer + Bible + diary', '❌ Separate apps'),
    ('Offline-first architecture', '✅ All features work offline', '❌ Most require internet'),
    ('Zero data collection', '✅ No accounts required, no tracking', '❌ Ad-based or data-harvesting'),
    ('Customizable', '✅ 5 themes, light/dark, fonts, layouts', '❌ Limited or premium-only'),
    ('Free & open source', '✅ Full source on GitHub', '⚠️ Often freemium'),
]

for diff, bf, comp in diffs:
    row = diff_table.add_row().cells
    row[0].text = diff
    row[1].text = bf
    row[2].text = comp

doc.add_page_break()

# ---- 5. Target Audience ----
doc.add_heading('5. Target Audience', level=1)

doc.add_heading('5.1 Primary Audience', level=2)
doc.add_paragraph(
    'Practicing Christians aged 16–55 who own a smartphone and want to integrate their faith '
    'into daily productivity.'
)

doc.add_heading('5.2 User Personas', level=2)

persona_table = doc.add_table(rows=1, cols=6)
persona_table.style = 'Light Grid Accent 1'
hdr = persona_table.rows[0].cells
hdr[0].text = 'Persona'
hdr[1].text = 'Age'
hdr[2].text = 'Occupation'
hdr[3].text = 'Tech Level'
hdr[4].text = 'Faith Stage'
hdr[5].text = 'Primary Need'

personas = [
    ('Samuel — Busy Believer', '32', 'Software Engineer', 'High', 'Mature', 'Integrate faith into busy schedule'),
    ('Grace — New Christian', '24', 'Graduate Student', 'Medium', 'New believer', 'Structured spiritual growth'),
    ('David — Ministry Leader', '45', 'Worship Pastor', 'Low-medium', 'Mature leader', 'Ministry organization'),
    ('Esther — Privacy-Conscious', '38', 'Freelance Writer', 'High', 'Mature', 'Complete data privacy'),
    ('Caleb — Digital Native Teen', '17', 'High School Student', 'Very High', 'Exploring', 'Engaging, customizable'),
]

for persona, age, occ, tech, faith, need in personas:
    row = persona_table.add_row().cells
    row[0].text = persona
    row[1].text = age
    row[2].text = occ
    row[3].text = tech
    row[4].text = faith
    row[5].text = need

doc.add_heading('5.3 Geographic Scope', level=2)
geo = [
    'Primary: English-speaking markets (US, UK, Canada, Australia, Nigeria, Philippines)',
    'Secondary: Spanish, French, German, Portuguese (language selector implemented)',
    'Future: More languages via community contributions',
]
for item in geo:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.text = item

doc.add_page_break()

# ---- 6. Implementation Status ----
doc.add_heading('6. Implementation Status', level=1)

doc.add_heading('6.1 Completed Features', level=2)

completed = [
    ('Task Management', 'Full CRUD, categories, filters, undo, completion tracking', 'v1.0'),
    ('Prayer Tracker', 'Daily logging, streak calculation, last 5 logs', 'v1.0'),
    ('Bible Reader', '66 books, 12 translations, offline cache, recent reads', 'v1.0'),
    ('Diary/Journal', 'CRUD, mood picker, undo, date/time display', 'v1.0'),
    ('AI Faith Assistant', 'GROQ/OpenAI/OpenRouter integration, chat history', 'v1.0'),
    ('Hymn Book', '1,001 hymns, search, favorites, audio (54 tunes)', 'v3.1'),
    ('Daily Devotional', '365 devotionals, navigation, font size', 'v3.1'),
    ('Settings', 'Themes, dark/light, font size, layout, backup/restore', 'v1.0'),
    ('User Authentication', 'Email + Google OAuth, JWT tokens', 'v4.1'),
    ('Cloud Sync', 'Last-write-wins, 12 data types, PostgreSQL', 'v4.1'),
    ('RAG Bible Search', 'Pinecone vector search, 54 verses, 1024-dim', 'v4.1'),
    ('Multi-LLM', 'GROQ, OpenAI, OpenRouter provider selection', 'v4.1'),
    ('Freemium Model', 'Free + Premium tiers, auth-gated features', 'v4.1'),
    ('Legal Framework', '14 documents, in-app acceptance, Settings tab', 'v4.2'),
    ('Security Audit', 'Credentials removed, consistent dimensions, JSON fix', 'v4.2'),
    ('Unit Tests', '61/61 passing (Vitest)', 'v4.1'),
]

status_table = doc.add_table(rows=1, cols=3)
status_table.style = 'Light Grid Accent 1'
hdr = status_table.rows[0].cells
hdr[0].text = 'Feature'
hdr[1].text = 'Description'
hdr[2].text = 'Version'

for feature, desc, ver in completed:
    row = status_table.add_row().cells
    row[0].text = feature
    row[1].text = desc
    row[2].text = ver

doc.add_heading('6.2 Legal & Compliance Documents', level=2)

legal_docs = [
    ('Privacy Policy', '1.0.0', 'Complete'),
    ('Terms of Service', '1.0.0', 'Complete'),
    ('Terms of Use', '1.0.0', 'Complete'),
    ('Community Guidelines', '1.0.0', 'Complete'),
    ('Cookie Policy', '1.0.0', 'Complete'),
    ('Content Moderation Policy', '1.0.0', 'Complete'),
    ('Acceptable Use Policy', '1.0.0', 'Complete'),
    ('Third-Party Services', '1.0.0', 'Complete'),
    ('Data Retention Policy', '1.0.0', 'Complete'),
    ('Incident Response Plan', '1.0.0', 'Complete'),
    ('Data Compliance', '1.0.0', 'Complete'),
    ('Compliance Checklist', '1.0.0', 'Complete'),
    ('Security Policy', '1.0.0', 'Complete'),
    ('Data Collection Disclosure', '1.0.0', 'Complete'),
]

legal_table = doc.add_table(rows=1, cols=3)
legal_table.style = 'Light Grid Accent 1'
hdr = legal_table.rows[0].cells
hdr[0].text = 'Document'
hdr[1].text = 'Version'
hdr[2].text = 'Status'

for doc_name, ver, status in legal_docs:
    row = legal_table.add_row().cells
    row[0].text = doc_name
    row[1].text = ver
    row[2].text = status

doc.add_page_break()

# ---- 7. Updates & Upgrades So Far ----
doc.add_heading('7. Updates & Upgrades So Far', level=1)

doc.add_heading('7.1 Version History', level=2)

versions = [
    ('v1.0', '2026-06-09', 'Initial release: Tasks, Prayer, Bible, Diary, AI Chat, Settings'),
    ('v3.1.0', '2026-06-10', 'Hymns (1,001), Daily Devotionals (365), 7-tab nav, AI endpoints'),
    ('v3.1.0b', '2026-06-11', 'PWA, draggable nav, error boundary, accessibility, code splitting'),
    ('v4.0.0', '2026-07-03', 'Risk Register, Legal & Compliance framework, Scalability Plan'),
    ('v4.1.0', '2026-07-04', 'Auth, Google OAuth, Cloud Sync, Pinecone RAG, Multi-LLM, Freemium'),
    ('v4.2.0', '2026-07-04', 'Legal framework (14 docs), in-app acceptance, Settings Legal tab, security fixes'),
]

ver_table = doc.add_table(rows=1, cols=3)
ver_table.style = 'Light Grid Accent 1'
hdr = ver_table.rows[0].cells
hdr[0].text = 'Version'
hdr[1].text = 'Date'
hdr[2].text = 'Changes'

for ver, date, changes in versions:
    row = ver_table.add_row().cells
    row[0].text = ver
    row[1].text = date
    row[2].text = changes

doc.add_heading('7.2 Key Technical Achievements', level=2)

achievements = [
    'Implemented full user authentication system with email + Google OAuth',
    'Built cloud sync with last-write-wins conflict resolution for 12 data types',
    'Integrated Pinecone vector database for AI-powered Bible verse search',
    'Added multi-LLM support (GROQ, OpenAI, OpenRouter) with per-request provider selection',
    'Created freemium model with premium feature gating',
    'Deployed PostgreSQL on Aiven for cloud data storage',
    'Implemented 61 unit tests with Vitest (all passing)',
    'Created comprehensive legal & compliance framework (14 documents)',
    'Built in-app legal acceptance flow with backend tracking',
    'Fixed security issues: removed hardcoded credentials, consistent embedding dimensions',
    'Improved UI: all tabs visible, settings navigation always accessible',
]

for achievement in achievements:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.text = achievement

doc.add_heading('7.3 Security Improvements', level=2)

security = [
    'Removed hardcoded database credentials from source code',
    'Removed hardcoded API keys (Pinecone, OpenAI) from source code',
    'Added JWT secret key warning for production environments',
    'Fixed embedding dimension mismatch (1024 consistent across all modules)',
    'Fixed JSON serialization in legal acceptance endpoint',
    'Added proper error handling for missing environment variables',
]

for item in security:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.text = item

doc.add_page_break()

# ---- 8. Expected Upgrades ----
doc.add_heading('8. Expected Upgrades', level=1)

doc.add_heading('8.1 v5.0 — Payment Integration & Group Features', level=2)

v5_features = [
    ('Flutterwave Payment Integration', 'Subscription billing, M-Pesa support, 15+ African countries', 'High'),
    ('Small Group Features', 'Create/join groups, shared prayer lists, invite codes', 'High'),
    ('Push Notifications', 'FCM integration, prayer reminders, verse-of-day', 'Medium'),
    ('Hebrew/Greek Interlinear', 'Original language Bible study tools', 'Medium'),
    ('Church Directory', 'Church profiles, leader management', 'Medium'),
    ('Advanced Analytics', 'Prayer streak analytics, reading patterns', 'Low'),
]

v5_table = doc.add_table(rows=1, cols=3)
v5_table.style = 'Light Grid Accent 1'
hdr = v5_table.rows[0].cells
hdr[0].text = 'Feature'
hdr[1].text = 'Description'
hdr[2].text = 'Priority'

for feature, desc, priority in v5_features:
    row = v5_table.add_row().cells
    row[0].text = feature
    row[1].text = desc
    row[2].text = priority

doc.add_heading('8.2 v5.1 — Enhanced AI & Community', level=2)

v51_features = [
    ('AI Bible Commentary', 'Theological commentary generation', 'High'),
    ('Community Forum', 'In-app discussion boards', 'Medium'),
    ('Gospel Music Streaming', 'Integration with music services', 'Medium'),
    ('Sermon Notes', 'Advanced sermon preparation tools', 'Medium'),
    ('Multi-language Support', 'i18n for Spanish, French, Portuguese', 'Low'),
]

v51_table = doc.add_table(rows=1, cols=3)
v51_table.style = 'Light Grid Accent 1'
hdr = v51_table.rows[0].cells
hdr[0].text = 'Feature'
hdr[1].text = 'Description'
hdr[2].text = 'Priority'

for feature, desc, priority in v51_features:
    row = v51_table.add_row().cells
    row[0].text = feature
    row[1].text = desc
    row[2].text = priority

doc.add_heading('8.3 v5.2 — Platform Expansion', level=2)

v52_features = [
    ('iOS App', 'Capacitor-based iOS build', 'High'),
    ('Desktop App', 'Electron-based desktop version', 'Medium'),
    ('Wearable Support', 'Apple Watch / WearOS companion', 'Low'),
    ('Offline RAG', 'Local vector search without cloud', 'Medium'),
]

v52_table = doc.add_table(rows=1, cols=3)
v52_table.style = 'Light Grid Accent 1'
hdr = v52_table.rows[0].cells
hdr[0].text = 'Feature'
hdr[1].text = 'Description'
hdr[2].text = 'Priority'

for feature, desc, priority in v52_features:
    row = v52_table.add_row().cells
    row[0].text = feature
    row[1].text = desc
    row[2].text = priority

doc.add_heading('8.4 Expected Upgrades Based on Recent Work', level=2)

doc.add_paragraph(
    'Based on the legal framework implementation and security audit completed in v4.2.0, '
    'the following upgrades are now enabled:'
)

expected = [
    ('Payment Processing', 'Legal framework enables subscription billing via Flutterwave'),
    ('User Trust', 'Transparent legal documents build user confidence for cloud sync'),
    ('Enterprise Adoption', 'Compliance framework enables church/ministry deployments'),
    ('International Expansion', 'Legal documents support GDPR/CCPA for global users'),
    ('Data Governance', 'Retention policies and security audits enable enterprise data management'),
    ('Community Features', 'Community Guidelines enable forum and group features'),
    ('API Partnerships', 'Third-Party Services disclosure enables future integrations'),
]

for feature, desc in expected:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    run = p.add_run(f'{feature}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ---- 9. Technical Architecture ----
doc.add_heading('9. Technical Architecture', level=1)

doc.add_heading('9.1 Stack Overview', level=2)

stack = [
    ('Frontend', 'React 19 + Vite 8 + Capacitor 8'),
    ('Backend', 'Python FastAPI + asyncpg'),
    ('Database', 'PostgreSQL on Aiven'),
    ('Vector Store', 'Pinecone (1024-dim embeddings)'),
    ('AI Providers', 'GROQ, OpenAI, OpenRouter'),
    ('Auth', 'JWT + Google OAuth'),
    ('Hosting', 'Vercel (serverless)'),
    ('Testing', 'Vitest (61 tests)'),
]

stack_table = doc.add_table(rows=1, cols=2)
stack_table.style = 'Light Grid Accent 1'
hdr = stack_table.rows[0].cells
hdr[0].text = 'Component'
hdr[1].text = 'Technology'

for comp, tech in stack:
    row = stack_table.add_row().cells
    row[0].text = comp
    row[1].text = tech

doc.add_heading('9.2 API Endpoints', level=2)

endpoints = [
    ('POST /api/auth/register', 'User registration'),
    ('POST /api/auth/login', 'User login'),
    ('POST /api/auth/google', 'Google OAuth'),
    ('POST /api/auth/legal-accept', 'Record legal acceptance'),
    ('GET /api/auth/legal-status', 'Check legal status'),
    ('GET /api/sync/pull', 'Pull user data from cloud'),
    ('POST /api/sync/push', 'Push user data to cloud'),
    ('POST /api/llm/chat', 'AI chat endpoint'),
    ('POST /api/rag/search', 'Bible verse search'),
    ('POST /api/rag/ingest', 'Ingest Bible verses'),
    ('GET /api/health', 'Health check'),
    ('GET /api/dbtest', 'Database test'),
    ('GET /api/pinetest', 'Pinecone test'),
]

ep_table = doc.add_table(rows=1, cols=2)
ep_table.style = 'Light Grid Accent 1'
hdr = ep_table.rows[0].cells
hdr[0].text = 'Endpoint'
hdr[1].text = 'Description'

for ep, desc in endpoints:
    row = ep_table.add_row().cells
    row[0].text = ep
    row[1].text = desc

doc.add_page_break()

# ---- 10. Legal & Compliance Framework ----
doc.add_heading('10. Legal & Compliance Framework', level=1)

doc.add_heading('10.1 Regulatory Compliance', level=2)

regulations = [
    ('GDPR (EU/EEA)', 'Partial', 'Privacy policy, user rights, security implemented'),
    ('UK GDPR', 'Partial', 'Same as GDPR implementation'),
    ('CCPA/CPRA (California)', 'Partial', 'User rights, no data selling'),
    ('NDPR (Nigeria)', 'Partial', 'Consent, security, breach notification'),
    ('LGPD (Brazil)', 'Partial', 'Similar to GDPR'),
    ('PIPEDA (Canada)', 'Partial', 'Consent, access rights'),
    ('POPIA (South Africa)', 'Partial', 'Similar to GDPR'),
]

reg_table = doc.add_table(rows=1, cols=3)
reg_table.style = 'Light Grid Accent 1'
hdr = reg_table.rows[0].cells
hdr[0].text = 'Regulation'
hdr[1].text = 'Status'
hdr[2].text = 'Implementation'

for reg, status, impl in regulations:
    row = reg_table.add_row().cells
    row[0].text = reg
    row[1].text = status
    row[2].text = impl

doc.add_heading('10.2 Items Requiring Legal Review', level=2)

review_items = [
    'Privacy Policy legal review by qualified counsel',
    'Terms of Service legal review by qualified counsel',
    'Terms of Use legal review by qualified counsel',
    'Data Protection Impact Assessment (DPIA)',
    'Data processing agreements with third-party providers',
    'Security audit by security professional',
    'Penetration testing by security professional',
]

for item in review_items:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.text = item

doc.add_page_break()

# ---- 11. Risk Assessment ----
doc.add_heading('11. Risk Assessment', level=1)

risks = [
    ('Security', 'High', 'Hardcoded credentials (fixed in v4.2.0)', 'Low', 'Credentials removed, env vars required'),
    ('Compliance', 'High', 'Legal documents not reviewed by counsel', 'Medium', 'Template created, counsel review pending'),
    ('Performance', 'Medium', 'Serverless cold starts', 'Low', 'Connection pooling implemented'),
    ('Scalability', 'Medium', 'User growth beyond free tier', 'Low', 'Infrastructure cost projections documented'),
    ('Dependency', 'Medium', 'Third-party API changes', 'Low', 'Multi-provider strategy (GROQ/OpenAI/OpenRouter)'),
    ('Data Loss', 'High', 'No cloud backup for local users', 'Medium', 'Export/import + cloud sync available'),
]

risk_table = doc.add_table(rows=1, cols=4)
risk_table.style = 'Light Grid Accent 1'
hdr = risk_table.rows[0].cells
hdr[0].text = 'Category'
hdr[1].text = 'Impact'
hdr[2].text = 'Risk'
hdr[3].text = 'Mitigation'

for cat, impact, risk, level, mitigation in risks:
    row = risk_table.add_row().cells
    row[0].text = cat
    row[1].text = impact
    row[2].text = risk
    row[3].text = mitigation

doc.add_page_break()

# ---- 12. Next Steps ----
doc.add_heading('12. Next Steps', level=1)

doc.add_heading('12.1 Immediate (Next 2 Weeks)', level=2)

immediate = [
    'Legal counsel review of all 14 legal documents',
    'Deploy frontend to Vercel with production backend URL',
    'iOS Capacitor build and testing',
    'Implement consent management UI for new users',
    'Build Privacy Dashboard in-app',
]

for item in immediate:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.text = item

doc.add_heading('12.2 Short-term (Next Month)', level=2)

short_term = [
    'Flutterwave payment integration',
    'Push notification system (FCM)',
    'Small group features',
    'Hebrew/Greek interlinear Bible',
    'Community forum',
]

for item in short_term:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.text = item

doc.add_heading('12.3 Medium-term (Next Quarter)', level=2)

medium_term = [
    'iOS App Store submission',
    'Church directory feature',
    'Advanced analytics dashboard',
    'Multi-language support (Spanish, French, Portuguese)',
    'Gospel music streaming integration',
]

for item in medium_term:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    p.text = item

doc.add_paragraph()
doc.add_paragraph()

# ---- Footer ----
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— End of Document —')
run.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ---- Save ----
output_path = r'C:\Users\ibrah\Documents\Gemini\Christian_Todo\pitch\mentorprd.docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")
