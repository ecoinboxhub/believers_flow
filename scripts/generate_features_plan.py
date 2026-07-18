from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def make_table(doc, headers, data):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for par in cell.paragraphs:
            for r in par.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row_idx, row_data in enumerate(data):
        for col_idx, val in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = val
            for par in cell.paragraphs:
                for r in par.runs:
                    r.font.size = Pt(9)

# ===== TITLE =====
title = doc.add_heading('BelieversFlow \u2014 Features & Implementation Plan', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Freemium Model Expansion \u2014 Phase 2 Planning')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x4A, 0x1A, 0x6B)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f'Generated: {datetime.datetime.now().strftime("%B %d, %Y")}\nStatus: Planning / Pre-Development')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

# ===== SECTION 1 =====
doc.add_heading('1. Current App Architecture (v3.1.0 Baseline)', level=1)
p = doc.add_paragraph()
p.add_run('Before implementing any new features, it is critical to understand the existing architecture because every proposed enhancement will require significant refactoring of the current codebase.').bold = True

make_table(doc,
    ['Layer', 'Technology', 'Limitation for Expansion'],
    [
        ['Frontend', 'React 19 + Vite 8, single-file SPA\n(App.jsx = ~1,932 lines)', 'Monolithic \u2014 must be refactored into modular components before adding features. No router, no DI.'],
        ['State Management', 'React useState + useCallback,\nlocalStorage (btf_* keys)', 'No global state store. Feature gating, auth, and cloud sync impossible without a proper state layer.'],
        ['Backend', 'FastAPI (Vercel serverless,\n358 lines, 12 endpoints)', 'Stateless, no database, no user accounts. Must add DB, auth, and payment processing.'],
        ['AI', 'GROQ llama-3.3-70b (generic\nmodel, no fine-tuning)', 'No RAG pipeline, no curated knowledge base, no fine-tuned model.'],
        ['Storage', 'Browser localStorage only', 'Not suitable for cloud sync, user accounts, or multi-device access.'],
        ['Auth', 'None (no user accounts)', 'Freemium requires user identification, authentication, and session management.'],
        ['Payments', 'None', 'No Stripe, no subscription management, no licensing.'],
        ['Content', 'Bundled JS files\n(hymns.js, devotional.js)', 'Static \u2014 cannot add multi-church content without a CMS or API-driven content delivery.'],
        ['Audio', 'Web Audio API synthesis\n(triangle wave, 54 tunes)', 'No streaming, no external music source integration.'],
        ['Community', 'None', 'No forum, no user profiles, no messaging. Needs a full social layer.'],
    ])
doc.add_paragraph()

# ===== SECTION 2 =====
doc.add_heading('2. Proposed Features \u2014 Detailed Breakdown', level=1)

# Feature 1
doc.add_heading('2.1 Multi-Church Devotional Content', level=2)
p = doc.add_paragraph()
p.add_run('Vision: ').bold = True
p.add_run('Provide daily devotionals from many churches and ministries \u2014 Christ Embassy, Winners Chapel, Dunamis, RCCG, Catholic Church, ECWA, Hillsong, and many more \u2014 all accessible within a single platform.')

p = doc.add_paragraph()
p.add_run('Suggested Church List (expandable): ').bold = True
churches = (
    'Christ Embassy (LoveWorld), Living Faith Church (Winners Chapel), Dunamis International Gospel Centre, '
    'RCCG (The Redeemed Christian Church of God), Catholic Church (daily readings), ECWA (Evangelical Church Winning All), '
    'Hillsong Church, Daystar Church, Joyce Meyer Ministries, Kenneth Copeland Ministries, '
    'Andrew Wommack Ministries, Joseph Prince Ministries, Charles Spurgeon (Morning & Evening), '
    'Our Daily Bread, Streams in the Desert, My Utmost for His Highest (Oswald Chambers), '
    'Bethlehem Baptist (John Piper), Desiring God, The Bible Project, '
    'Alpha International, HTB (Holy Trinity Brompton), '
    'Soul Food (Catholic), Word Among Us (Catholic), '
    'Greek Orthodox Daily Readings, Coptic Orthodox Readings, '
    'African Gospel Church, Deeper Life Bible Church, '
    'Salvation Army, Seventh-day Adventist (devotional), '
    'Assemblies of God, Baptist General Convention, '
    'Anglican Communion Daily Prayer, Presbyterian Church (PCUSA), '
    'Methodist Church Daily Readings, Lutheran Daily Devotion, '
    'Mennonite Devotional, Brethren Daily Readings, '
    'Calvary Chapel, Vineyard Churches, '
    'Sovereign Grace Churches, Acts 29 Network'
)
p.add_run(churches)

steps = [
    ('Content Partnerships', 'Reach out to each church/ministry for permission to redistribute daily devotionals. Formal licensing agreements required.'),
    ('Content Management System', 'Build a CMS (admin panel) where churches upload daily devotionals (title, verse, reflection, prayer). Each church gets a dashboard.'),
    ('API-Driven Delivery', 'Devotionals served via API (not bundled). Users subscribe to churches they follow. Free tier: 1 church. Premium: unlimited.'),
    ('Database Schema', 'PostgreSQL/Supabase tables: churches, devotional_entries, user_church_subscriptions, read_history.'),
    ('Caching Strategy', 'Cache today\'s devotionals in localStorage for offline reading. Fetch new content on app open.'),
]
for title, desc in steps:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{title}: ').bold = True
    p.add_run(desc)

# Feature 2
doc.add_heading('2.2 AI Assistant Trained on Trusted Biblical Sources', level=2)
p = doc.add_paragraph()
p.add_run('Vision: ').bold = True
p.add_run('Replace the current generic GROQ model with an AI that answers exclusively from curated biblical and theological sources \u2014 the Bible, trusted commentaries, and orthodox Christian literature.')

steps = [
    ('Source Curation', 'Curate a corpus: multiple Bible translations (KJV, NIV, ESV, NASB), classic commentaries (Matthew Henry, John Gill, Jamieson-Fausset-Brown, Cambridge Bible), systematic theologies (Grudem, Berkhof), and public-domain Christian classics (Augustine, Luther, Wesley, Spurgeon).'),
    ('RAG Pipeline (Retrieval-Augmented Generation)', 'Chunk all documents, embed with a sentence-transformer (e.g., all-MiniLM-L6-v2), store in a vector database (pgvector or Chroma). On each user query, retrieve top-k relevant passages and inject into the LLM prompt as context.'),
    ('Fine-Tuning (Optional Phase 2)', 'Create a high-quality Q&A dataset from curated sources. Fine-tune an open-source model (Llama 3 8B, Mixtral 8x7B) on this data for more theologically aligned responses.'),
    ('Guardrails', 'Implement output filtering to ensure responses cite specific sources. If no source matches, AI responds: "I don\'t have a trusted source for that." No speculative theology.'),
    ('Hosting', 'RAG pipeline can run on the same FastAPI backend. Vector DB via Supabase (pgvector) in Premium tier. Free tier: limited to GROQ prompt engineering only.'),
]
for title, desc in steps:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{title}: ').bold = True
    p.add_run(desc)

# Feature 3
doc.add_heading('2.3 Gospel Music Integration', level=2)
p = doc.add_paragraph()
p.add_run('Vision: ').bold = True
p.add_run('Move beyond the current 54 synthesized hymn tunes to a full gospel music experience \u2014 streaming worship songs, praise music, and hymns from various Christian traditions.')

steps = [
    ('Music Source Options', 'Option A: Integrate with a Christian music API (e.g., Spotify API for gospel playlists, Apple Music, or a dedicated Christian music provider). Option B: Host music files on CDN (requires licensing). Option C: Curate YouTube/embed links to legal worship content.'),
    ('Licensing & Royalties', 'Critical path \u2014 must secure licenses (CCLI, SongSelect, or direct publisher agreements). Cannot host copyrighted music without permission. Consider partnerships with gospel labels.'),
    ('Audio Player', 'Replace the current Web Audio synthesis player with an HTML5 <audio> or Web Audio streaming player. Features: play/pause, skip, seek, playlist, offline download (Premium).'),
    ('Music Catalog', 'Organize by genre (worship, praise, hymns, gospel, contemporary), artist, album, and theme (Christmas, Easter, Communion).'),
    ('Offline Listening', 'Free tier: streaming only. Premium tier: download songs for offline listening.'),
]
for title, desc in steps:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{title}: ').bold = True
    p.add_run(desc)

# Feature 4
doc.add_heading('2.4 Community Forum & Fellowship Platform', level=2)
p = doc.add_paragraph()
p.add_run('Vision: ').bold = True
p.add_run('Create a space where believers from different churches can engage, share insights, discuss biblical topics, ask questions, and build meaningful connections.')

steps = [
    ('Architecture Decision', 'Option A: Build a custom forum within the app (React + backend API + DB). Option B: Embed an existing forum solution (Discourse, Flarum, TalkYard) via iframe or API. Recommended: Option A for seamless UX, Option B as MVP shortcut.'),
    ('Forum Features', 'Category-based discussions (Bible Study, Prayer Requests, Testimonies, Theology Q&A, Church Announcements). Threaded replies, upvoting, moderation tools, user reputation.'),
    ('User Profiles', 'Public profiles with church affiliation, bio, and activity history. Built on the same auth system used for freemium.'),
    ('Moderation', 'AI-assisted moderation (flag offensive content, detect spam). Human moderators from each church community. Reporting system.'),
    ('Real-Time', 'WebSocket (Socket.IO) for real-time updates on new posts, replies, and notifications.'),
]
for title, desc in steps:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{title}: ').bold = True
    p.add_run(desc)

# Feature 5
doc.add_heading('2.5 Freemium Model & Modular Feature Selection', level=2)
p = doc.add_paragraph()
p.add_run('Vision: ').bold = True
p.add_run('Users can choose which features they want enabled in their app, and the freemium model gates premium features behind a subscription, while keeping core spiritual tools free forever.')

steps = [
    ('User Accounts & Auth', 'Add authentication (email/password + Google SSO). Use Supabase Auth or Firebase Auth. Required for subscription management and cloud sync.'),
    ('Subscription Management', 'Stripe for payment processing. Plans: Free ($0), Supporter ($2.99/mo), Ministry ($9.99/mo). Annual discount (20% off). Webhook-based license key management.'),
    ('Feature Gating', 'Create a FeatureFlag system: premium features are toggled server-side based on subscription tier. The frontend checks flags on init and conditionally renders features.'),
    ('Modular UI Architecture', 'Refactor the monolithic App.jsx into independently loadable modules (plugins). Users see a "Feature Store" where they enable/disable modules. Disabled modules are lazy-loaded (code-split) and never executed.'),
    ('Free Tier Boundaries', 'Free: 1 church devotional subscription, 30 AI messages/day, streaming-only music, read-only forum. Premium: unlimited churches, unlimited AI (RAG-powered), offline music downloads, full forum participation, cloud sync, no branding.'),
    ('Local-First Sync', 'All features work offline first. Cloud sync happens in background when online. Conflicts resolved with last-write-wins or user choice.'),
]
for title, desc in steps:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{title}: ').bold = True
    p.add_run(desc)

doc.add_paragraph()

# ===== SECTION 3 =====
doc.add_heading('3. Technical Architecture Recommendations', level=1)

doc.add_heading('3.1 Frontend Refactoring', level=2)
for item in [
    'Decompose App.jsx (~1,932 lines) into separate component files per view (TaskManager.jsx, BibleReader.jsx, HymnBook.jsx, etc.).',
    'Adopt a lightweight state management solution (Zustand) for global state (auth, features, settings).',
    'Introduce a client-side router (React Router) for proper URL-based navigation and deep linking.',
    'Implement code-splitting per module \u2014 users only download features they have enabled.',
    'Add a plugin/module registry system where features register themselves (icon, route, premium status).',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.2 Backend Expansion', level=2)
for item in [
    'Add a database: Supabase (PostgreSQL) recommended \u2014 offers auth, DB, storage, real-time, and vector search in one platform. Free tier is generous.',
    'Add user authentication: Supabase Auth handles email/password, Google SSO, session management.',
    'Add payment processing: Stripe Checkout for subscriptions. Webhooks to update user tiers in Supabase.',
    'RESTful API expansion: New endpoints for devotionals (CRUD per church), music catalog, forum posts, user profiles, subscriptions.',
    'Consider a dedicated backend server (Node.js/Express or Python FastAPI on a VPS or Railway) instead of Vercel serverless, since you\'ll need persistent connections (WebSockets) and background jobs.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.3 AI Pipeline', level=2)
for item in [
    'Phase 1: Improve GROQ prompts with curated source citations (no RAG yet). Add "Based on [source]" footnotes.',
    'Phase 2: Build RAG pipeline \u2014 vector embed curated documents, store in pgvector, retrieve on each query.',
    'Phase 3: Fine-tune an open-source model on curated Q&A dataset. Deploy on Groq, Together, or self-host via vLLM.',
    'Maintain a source attribution system \u2014 every AI response must cite its source document and page/verse.',
    'Premium tier exclusive: RAG-powered AI with full source attribution. Free tier gets prompt-engineered GROQ only.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3.4 Data & Content Pipeline', level=2)
for item in [
    'Content Management Dashboard: Build a web admin panel where churches submit devotionals (approval workflow).',
    'CDN for static assets: Church devotionals, music files, and images served via CDN (Vercel Edge, Cloudflare R2, or Supabase Storage).',
    'Offline-first sync engine: Queue user actions (write to localStorage immediately, sync to server when online). Use a sync log (similar to CRDT or OT).',
    'Content licensing: Maintain a licensing database per church/content source to track rights and expiration.',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

# ===== SECTION 4 =====
doc.add_heading('4. Phased Implementation Roadmap', level=1)

p = doc.add_paragraph()
p.add_run('The planning & preparation phase (no coding) is compressed into a single month. The actual coding phases follow.').bold = True

doc.add_heading('Month 0 \u2014 Planning & Preparation (No Coding)', level=2)
make_table(doc,
    ['Week', 'Focus', 'Deliverables'],
    [
        ['Week 1', 'Content Partnership Outreach + User Survey', 'Draft church proposal letters. Identify 3-5 pilot churches. Survey current users on willingness to pay and desired features.'],
        ['Week 1-2', 'Architecture Review & Tech Stack Decision', 'Finalize state management (Zustand), backend stack (Node.js vs FastAPI), database (Supabase vs Firebase), payment processor (Stripe). Create detailed migration plan for decomposing App.jsx.'],
        ['Week 2-3', 'Music Licensing + AI Source Curation', 'Contact CCLI/SongSelect for pricing. Compile list of public-domain commentaries and theological works. Check copyright status. Organize by Bible book and topic.'],
        ['Week 3', 'Wireframing & UX Design', 'Design "Feature Store" UI, church subscription picker, forum layout, music player, premium upgrade prompts. Create Figma mockups for all new screens.'],
        ['Week 3-4', 'Legal + Compliance + Financial Model', 'Terms of service, privacy policy, subscription agreement, DMCA compliance, content licensing contracts. Refine pricing model. Set break-even target.'],
    ])
doc.add_paragraph()

doc.add_heading('Phase 1 \u2014 Foundation Coding (Months 2-3)', level=2)
make_table(doc,
    ['Month', 'Focus', 'Deliverables'],
    [
        ['Month 2', 'Architecture Refactor', 'Decompose App.jsx into modules. Add Zustand state management. Set up React Router. Create plugin/module registry system.'],
        ['Month 2-3', 'Backend & Auth', 'Provision Supabase project. Implement user auth (email + Google SSO). Create user profiles DB. Port backend to Node.js or expand FastAPI.'],
        ['Month 3', 'Feature Gating + Payments', 'Build FeatureFlag system. Stripe Checkout integration. Subscription webhook handling. Modular feature selection UI ("Feature Store").'],
    ])
doc.add_paragraph()

doc.add_heading('Phase 2 \u2014 Core Features Coding (Months 4-6)', level=2)
make_table(doc,
    ['Month', 'Focus', 'Deliverables'],
    [
        ['Month 4', 'Cloud Sync + Multi-Church Devotionals', 'Background sync engine. CMS admin panel. Onboard 3-5 pilot churches. Devotional API with church subscription system.'],
        ['Month 4-5', 'AI RAG Pipeline', 'Curate source corpus. Build vector embeddings. Implement RAG retrieval. Integrate with chat UI. Premium gating.'],
        ['Month 5', 'Gospel Music MVP', 'Secure 1-2 licensing partnerships. Implement audio streaming player. Catalog 50-100 songs. Offline download for Premium.'],
        ['Month 5-6', 'Community Forum MVP + Polish', 'Thread/reply system with real-time (WebSockets). User profiles with church affiliation. Basic moderation. Push notifications.'],
    ])
doc.add_paragraph()

doc.add_heading('Phase 3 \u2014 Scale & Distribute (Months 7-9)', level=2)
make_table(doc,
    ['Month', 'Focus', 'Deliverables'],
    [
        ['Month 7', 'Content Expansion', 'Onboard 20+ churches. Expand music catalog to 500+ songs. AI fine-tuning on curated data.'],
        ['Month 7-8', 'AI Moderation + Analytics', 'AI-assisted forum moderation. Prayer analytics dashboard. White-label branding option.'],
        ['Month 8-9', 'iOS & Marketing', 'Build iOS version (Capacitor on macOS). Submit to App Store. Marketing campaign.'],
    ])
doc.add_paragraph()

# ===== SECTION 5 =====
doc.add_heading('5. Cost & Resource Estimates', level=1)

make_table(doc,
    ['Item', 'Monthly Cost (Start)', 'Monthly Cost (Scale)', 'Notes'],
    [
        ['Supabase (Pro)', '$25', '$50', 'DB, Auth, Storage, Realtime, pgvector'],
        ['Backend Hosting', '$0 (Vercel)', '$20-50 (Railway/VPS)', 'Node.js or FastAPI server'],
        ['Stripe', '$0', '2.9% + $0.30/txn', 'Payment processing fees'],
        ['GROQ API', '$0 (free tier)', '$50-200', 'Prompt-engineered AI for free tier'],
        ['Vector Embeddings + RAG', '$0', '$50-100', 'OpenAI or self-hosted embedding'],
        ['CDN / Storage', '$0 (Supabase)', '$10-30', 'Music files, images, content assets'],
        ['Music Licensing', '$0 (negotiate)', '$200-1000', 'CCLI license or direct publisher deals'],
        ['AI Fine-Tuning (one-time)', '$0', '$500-2000', 'Compute for fine-tuning (RunPod, Together)'],
        ['Total', '~$25/mo', '~$400-1500/mo', 'Revenue target: $0.50-1.00 per user/mo'],
    ])
doc.add_paragraph()

# ===== SECTION 6 =====
doc.add_heading('6. Critical Risks & Mitigation', level=1)

make_table(doc,
    ['Risk', 'Impact', 'Mitigation'],
    [
        ['Content licensing delays', 'High \u2014 blocks multi-church devotionals', 'Start conversations with churches early. Offer free promotion in exchange for content. Use public-domain devotionals as fallback.'],
        ['Music licensing costs too high', 'Medium \u2014 blocks gospel music', 'Start with royalty-free hymns and public-domain gospel songs. Partner with independent worship artists.'],
        ['Refactoring breaks existing features', 'High \u2014 user trust', 'Comprehensive test suite before refactoring. Feature flags to roll back. Gradual migration (strangler pattern).'],
        ['Low premium conversion (<2%)', 'High \u2014 revenue insufficient', 'A/B test pricing. Bundle premium features carefully. Add genuinely valuable premium-only features (cloud sync, offline music).'],
        ['4-week planning sprint too aggressive', 'Medium \u2014 incomplete preparation', 'Prioritize Week 1-2 as non-negotiable. Weeks 3-4 can spill into Month 2 if needed. Parallelize independent workstreams.'],
        ['Theological AI controversy', 'Medium \u2014 reputational risk', 'AI responds only from curated sources. Clear disclaimer: "AI is a tool, not a pastor." Human oversight for sensitive queries.'],
        ['Moderation costs for forum', 'Medium', 'AI-assisted moderation (automate 80%). Volunteer community moderators. Clear code of conduct.'],
        ['Supabase vendor lock-in', 'Low-Medium', 'Use standard PostgreSQL. Keep migration option open. Avoid Supabase-specific features where possible.'],
    ])
doc.add_paragraph()

# ===== SECTION 7 =====
doc.add_heading('7. Key Architectural Principles Going Forward', level=1)

principles = [
    ('Offline-First Always', 'All features must work without internet. Cloud sync is additive, not required. localStorage remains the source of truth; server is the backup.'),
    ('Modular by Design', 'Each feature is a self-contained module with its own state, views, and data layer. Modules can be enabled/disabled independently. No cross-module coupling.'),
    ('Privacy by Default', 'No telemetry, no analytics, no tracking. User data stays on device unless explicitly synced. End-to-end encryption for cloud-synced data.'),
    ('Denomination-Agnostic', 'The platform serves all Christian traditions. Content from different churches is clearly labeled. The AI does not take doctrinal positions.'),
    ('Free Core Forever', 'Bible reading, prayer tracking, journaling, and task management remain free indefinitely. Premium enhances but never withholds spiritual tools.'),
    ('Source Transparency', 'Every AI response cites its sources. Every devotional shows its church of origin. Every Bible verse shows its translation.'),
    ('Progressive Enhancement', 'Start with the current codebase. Add new features alongside existing ones. Migrate incrementally. Never rewrite from scratch.'),
]
for title, desc in principles:
    p = doc.add_paragraph()
    p.add_run(f'{title}: ').bold = True
    p.add_run(desc)

doc.add_paragraph()

# ===== SECTION 8 =====
doc.add_heading('8. One-Month Sprint: Planning & Preparation (No Coding)', level=1)

p = doc.add_paragraph()
p.add_run('All pre-coding work compressed into 4 weeks. Parallelize where possible.').bold = True

doc.add_paragraph()
doc.add_heading('Week 1: Foundation', level=2)
for title, desc in [
    ('Content Partnership Outreach', 'Draft proposal letters. Identify and contact 3-5 pilot churches (Christ Embassy, Winners Chapel, RCCG, Dunamis, Catholic Church). Secure content permission.'),
    ('User Survey', 'Survey current users via GitHub + in-app prompt. Validate willingness to pay, desired features, church affiliation. Target 100+ responses.'),
    ('Tech Stack Finalization', 'Decide: Zustand vs Jotai for state, Node.js vs FastAPI for backend, Supabase vs Firebase, Stripe vs Paddle. Document rationale in ADR (Architecture Decision Record).'),
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{title}: ').bold = True
    p.add_run(desc)

doc.add_heading('Week 2: Architecture & Design', level=2)
for title, desc in [
    ('App.jsx Decomposition Plan', 'Map every view/component in the current 1,932-line App.jsx. Design the module split: TaskManager, BibleReader, HymnBook, DevotionalView, DiaryView, SettingsPanel, SpiritualDashboard, AIChat. Plan data flow for each.'),
    ('Wireframing (Figma)', 'Design Feature Store UI (modular toggle screen), church subscription picker, community forum layout, gospel music player, premium upgrade prompts. Create clickable prototype.'),
    ('Database Schema Design', 'Design all new tables: users, profiles, subscriptions, churches, devotional_entries, user_church_subs, forum_threads, forum_posts, music_tracks, user_sync_log.'),
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{title}: ').bold = True
    p.add_run(desc)

doc.add_heading('Week 3: Content & Licensing', level=2)
for title, desc in [
    ('AI Source Curation', 'Compile list of public-domain commentaries (Matthew Henry, John Gill, Jamieson-Fausset-Brown, Cambridge Bible, Spurgeon, Augustine, Luther, Wesley). Download and organize by Bible book. Verify copyright status.'),
    ('Music Licensing Initiation', 'Contact CCLI for streaming license pricing. Research SongSelect API. Identify 3-5 independent gospel artists for pilot catalog. Evaluate royalty-free hymn repositories.'),
    ('Legal Drafting', 'Terms of service, privacy policy, subscription/license agreement, DMCA takedown policy, church content licensing template.'),
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{title}: ').bold = True
    p.add_run(desc)

doc.add_heading('Week 4: Finalization', level=2)
for title, desc in [
    ('Financial Model', 'Finalize pricing tiers ($0 / $2.99 / $9.99). Calculate break-even based on infrastructure costs. Set conversion targets (5% free-to-paid, 10% supporter-to-ministry).'),
    ('Implementation Backlog', 'Convert all planning into a prioritized GitHub Project board with user stories, acceptance criteria, and effort estimates. Tag by phase (Foundation, Core Features, Scale).'),
    ('Developer Environment Setup', 'Provision Supabase project. Configure Stripe test mode. Set up CI/CD for new backend. Create dev/staging branches. Document onboarding instructions in AGENTS.md.'),
    ('Go / No-Go Decision', 'Review all Week 1-4 deliverables. If content partnerships secured, licensing viable, and architecture plan solid -> proceed to Month 2 (coding phase).'),
]:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(f'{title}: ').bold = True
    p.add_run(desc)

doc.add_paragraph()

# ===== SECTION 9 =====
doc.add_heading('9. Freemium Tier Breakdown (Proposed)', level=1)

make_table(doc,
    ['Feature', 'Free', 'Supporter ($2.99/mo)', 'Ministry ($9.99/mo)'],
    [
        ['Task Management', '\u2705 Full', '\u2705 Full', '\u2705 Full'],
        ['Prayer Tracker', '\u2705 Full', '\u2705 Full', '\u2705 Full'],
        ['Bible Reader (66 books)', '\u2705 Full', '\u2705 Full', '\u2705 Full'],
        ['Diary / Journal', '\u2705 Full', '\u2705 Full', '\u2705 Full'],
        ['Daily Devotional (built-in)', '\u2705 Full', '\u2705 Full', '\u2705 Full'],
        ['Hymn Book (1,001)', '\u2705 Full', '\u2705 Full', '\u2705 Full'],
        ['Hymn Music (54 tunes)', '\u2705 Full', '\u2705 Full', '\u2705 Full'],
        ['AI Faith Assistant', '30 msg/day (prompt-only)', 'Unlimited (RAG-powered)', 'Unlimited (RAG-powered)'],
        ['Multi-Church Devotionals', '1 church', 'Up to 5 churches', 'Unlimited'],
        ['Church Content Dashboard', '\u274c', '\u274c', '\u2705 (admin panel)'],
        ['Gospel Music Streaming', 'Audio ads / limited', 'Unlimited streaming', 'Unlimited streaming'],
        ['Offline Music Downloads', '\u274c', '\u2705', '\u2705'],
        ['Community Forum', 'Read-only', 'Full participation', 'Full + priority'],
        ['Cloud Sync', '\u274c', '\u2705 (2 devices)', '\u2705 (unlimited)'],
        ['Branding Footer', 'Visible', 'Removed', 'White-label option'],
        ['Group Management', '\u274c', '\u274c', '\u2705 (up to 50 members)'],
        ['AI Moderation (Forum)', '\u274c', '\u274c', '\u2705'],
    ])
doc.add_paragraph()

# ===== CLOSING =====
doc.add_heading('Closing Note', level=1)
p = doc.add_paragraph()
p.add_run(
    'This document is a strategic planning guide \u2014 not a technical specification. '
    'Every recommendation here should be validated through user research, cost analysis, '
    'and technical prototyping before committing to implementation. The goal is to transform '
    'BelieversFlow from a personal productivity tool into a comprehensive Christian lifestyle '
    'platform, while preserving the privacy, offline-first, and free-core values that make it unique.'
)

p = doc.add_paragraph()
run = p.add_run('\n"The steadfast love of the Lord never ceases; his mercies never come to an end; they are new every morning." \u2014 Lamentations 3:22-23')
run.italic = True

output_path = r'C:\Users\ibrah\Documents\Gemini\Christian_Todo\pitch\features_plan.docx'
doc.save(output_path)
print(f"Document saved to {output_path}")
