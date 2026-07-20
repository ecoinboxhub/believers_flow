#!/usr/bin/env python3
"""Generate comprehensive BelieversFlow Implementation Audit Report"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0x1a, 0x0a, 0x2e)
    hs.font.name = 'Calibri'

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for row_data in rows:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = str(val)
    return t

def add_status_table(rows):
    """rows: list of (module, status_emoji, completion, priority, notes)"""
    return add_table(['Module', 'Status', 'Completion', 'Priority', 'Notes'], rows)

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(4):
    doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('BelieversFlow')
r.font.size = Pt(36)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1a, 0x0a, 0x2e)

st = doc.add_paragraph()
st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run('Comprehensive Implementation Audit Report')
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x7b, 0x2d, 0x8e)

doc.add_paragraph()
v = doc.add_paragraph()
v.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = v.add_run('PRD v4.2.0 — Production Hardening Audit')
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0xf2, 0xc9, 0x4c)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run(f'Audit Date: {datetime.date.today().strftime("%B %d, %Y")}\nClassification: Confidential — Internal Use Only')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc = [
    '1. Executive Audit Summary',
    '2. Production Readiness Scorecard',
    '3. PRD Compliance Report (Feature-by-Feature)',
    '4. Missing Features Report',
    '5. Security Audit',
    '6. Database Audit',
    '7. API Audit',
    '8. Architecture Review',
    '9. Code Quality Report',
    '10. Performance Review',
    '11. Technical Debt Report',
    '12. Remaining Implementation Checklist',
    '13. Prioritized Sprint Plan',
    '14. GitHub Issues / Tasks',
    '15. Final Recommendation',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE AUDIT SUMMARY
# ============================================================
doc.add_heading('1. Executive Audit Summary', level=1)

doc.add_paragraph(
    'This audit was conducted by inspecting the actual BelieversFlow codebase against the Product '
    'Requirements Document (PRD) v4.2.0. Every finding is evidence-based, referencing specific files, '
    'line numbers, and code. No feature is assumed to exist without verification.'
)

doc.add_heading('Overall Assessment', level=2)
doc.add_paragraph(
    'BelieversFlow is a functional MVP with core features implemented. However, it has significant '
    'gaps in security, testing, CI/CD, and several PRD-mandated features. The application is '
    'suitable for beta testing but NOT ready for production deployment without addressing critical '
    'security vulnerabilities and completing missing infrastructure.'
)

doc.add_heading('Key Metrics', level=2)
metrics = [
    ('Total Source Files (Frontend)', '11 JSX/JS files'),
    ('Total Source Files (Backend)', '8 Python files'),
    ('Total Lines of Code (Frontend)', '~4,500 lines (App.jsx alone: 2,278)'),
    ('Total Lines of Code (Backend)', '~1,400 lines'),
    ('Database Tables', '2 (users, user_data)'),
    ('API Endpoints', '24'),
    ('Unit Tests', '61 passing (2 test files)'),
    ('Legal Documents', '14'),
    ('Security Issues Found', '7 critical/high'),
    ('PRD Requirements Met', '~55% fully, ~20% partial, ~25% missing'),
    ('Production Readiness', '38/100'),
]
add_table(['Metric', 'Value'], metrics)

doc.add_heading('Critical Findings Summary', level=2)
critical = [
    ('CRITICAL', 'Hardcoded API keys in reingest.py (Pinecone + OpenAI)'),
    ('CRITICAL', 'Hardcoded JWT secret fallback in auth.py'),
    ('CRITICAL', 'No auth on RAG/LLM endpoints (cost abuse vector)'),
    ('HIGH', 'CORS allow_origins=["*"] with credentials'),
    ('HIGH', 'get_connection() bypasses connection pool (leak risk)'),
    ('HIGH', 'No password reset flow'),
    ('HIGH', 'No email verification'),
    ('HIGH', 'No account deletion endpoint'),
    ('HIGH', 'README.md is default Vite template, not project-specific'),
]
add_table(['Severity', 'Finding'], critical)

doc.add_page_break()

# ============================================================
# 2. PRODUCTION READINESS SCORECARD
# ============================================================
doc.add_heading('2. Production Readiness Scorecard', level=1)

doc.add_paragraph(
    'Score: 38/100 — Not production-ready. Requires significant work in security, testing, '
    'infrastructure, and missing features before production deployment.'
)

scorecard = [
    ('User Authentication', 'Partial', '45%', 'Critical', 'Email+Google works; no password reset, no email verification, no account deletion'),
    ('Database', 'Partial', '40%', 'Critical', '2 tables only; missing churches, ministries, subscriptions, feature flags, notifications'),
    ('Multi-Church Support', 'Missing', '0%', 'High', 'No church model, no ministry model, no publishing, no multi-tenant'),
    ('Offline-first', 'Partial', '65%', 'High', 'localStorage works; no offline queue, no delta sync, no conflict UI'),
    ('Cloud Sync', 'Partial', '50%', 'High', 'Basic push/pull works; no retry logic, no background sync, no conflict resolution UI'),
    ('Bible Module', 'Partial', '70%', 'Medium', '12 translations, search, notes; missing cross-references, verse sharing, highlighting'),
    ('AI Assistant', 'Partial', '60%', 'Medium', 'Multi-LLM works; missing source citations, confidence scores, rate limiting'),
    ('Prayer Module', 'Partial', '55%', 'Medium', 'Basic logging; missing categories, answer tracking, statistics'),
    ('Journal', 'Partial', '50%', 'Medium', 'CRUD works; missing encryption, markdown, search, export'),
    ('Hymn Module', 'Partial', '75%', 'Low', '1,001 hymns, audio, favorites; missing offline playback'),
    ('Personalization', 'Partial', '60%', 'Medium', 'Themes, nav reorder; missing hide features, denomination selection'),
    ('Feature Store', 'Missing', '0%', 'High', 'No feature flags, no remote config, no dynamic modules'),
    ('Feedback System', 'Missing', '0%', 'Medium', 'No bug reports, no ratings, no NPS, no surveys'),
    ('Subscription System', 'Missing', '0%', 'High', 'No billing, no purchase validation, no restore purchases'),
    ('Community', 'Missing', '0%', 'Medium', 'No forum, no posts, no moderation'),
    ('Security', 'Weak', '35%', 'Critical', 'Hardcoded secrets, open CORS, no rate limiting, no RBAC'),
    ('Privacy & Compliance', 'Partial', '55%', 'High', '14 legal docs exist; no consent UI, no privacy dashboard, no DPIA'),
    ('Notifications', 'Missing', '0%', 'Medium', 'No push, no local, no scheduled reminders'),
    ('Performance', 'Partial', '50%', 'Medium', 'PWA works; no lazy loading, 1MB+ bundle, no image optimization'),
    ('Testing', 'Weak', '25%', 'Critical', '61 unit tests; no integration tests, no E2E tests, no backend tests'),
    ('Documentation', 'Partial', '40%', 'Medium', 'Legal docs good; README is default Vite, no API docs, no architecture docs'),
    ('CI/CD', 'Weak', '20%', 'High', 'Only backend deploy workflow; no frontend CI, no lint, no test pipeline'),
]

add_status_table(scorecard)

doc.add_page_break()

# ============================================================
# 3. PRD COMPLIANCE REPORT
# ============================================================
doc.add_heading('3. PRD Compliance Report (Feature-by-Feature)', level=1)

doc.add_heading('3.1 Executive Vision', level=2)
vision_items = [
    ('Offline-first architecture', 'Partial', 'localStorage caching works; no offline queue, no delta sync', 'App.jsx:304-335'),
    ('Productivity + Faith integration', 'Full', 'Tasks, prayer, Bible, diary, hymns, devotionals integrated', 'App.jsx:1037-1795'),
    ('Multi-platform support', 'Partial', 'PWA works; Android Capacitor configured but no iOS', 'capacitor.config.json'),
    ('PWA capabilities', 'Full', 'Service worker, manifest, caching configured', 'vite.config.js:8-52'),
    ('Android readiness', 'Partial', 'Capacitor config exists; no build automation', 'capacitor.config.json'),
    ('iOS readiness', 'Missing', 'No capacitor.config.ts for iOS, no iOS build', 'N/A'),
]
add_table(['Requirement', 'Status', 'Evidence', 'Location'], vision_items)

doc.add_heading('3.2 User Authentication', level=2)
auth_items = [
    ('Email signup', 'Full', 'POST /api/auth/register creates user', 'auth.py:100-115'),
    ('Email login', 'Full', 'POST /api/auth/login validates credentials', 'auth.py:117-131'),
    ('Logout', 'Full', 'Frontend clears tokens, stops sync', 'syncService.js:27-31'),
    ('Password reset', 'Missing', 'No endpoint, no UI, no email service', 'N/A'),
    ('Email verification', 'Missing', 'No verification flow', 'N/A'),
    ('Google OAuth', 'Full', 'POST /api/auth/google with token verification', 'auth.py:133-180'),
    ('Secure sessions (JWT)', 'Partial', 'JWT with 30-day expiry; no refresh tokens', 'auth.py:64-69'),
    ('Refresh tokens', 'Missing', 'No refresh token mechanism', 'N/A'),
    ('Role-based permissions', 'Missing', 'No roles, no RBAC, no admin', 'N/A'),
    ('Guest mode', 'Full', 'Local-only mode without account', 'App.jsx:983-985'),
    ('Account deletion', 'Missing', 'No endpoint to delete user account', 'N/A'),
]
add_table(['Requirement', 'Status', 'Evidence', 'Location'], auth_items)

doc.add_heading('3.3 Database Schema', level=2)
db_items = [
    ('Users table', 'Full', 'id, email, name, password_hash, plan, created_at', 'database.py:37-44'),
    ('User data table', 'Full', 'id, user_id, data_type, data (JSONB), updated_at', 'database.py:53-62'),
    ('Churches table', 'Missing', 'No church model exists', 'N/A'),
    ('Ministries table', 'Missing', 'No ministry model exists', 'N/A'),
    ('Devotionals table', 'Missing', 'Hardcoded in JS, not in DB', 'devotional.js'),
    ('Notes table', 'Missing', 'No notes table', 'N/A'),
    ('Journals table', 'Missing', 'Stored in localStorage only', 'N/A'),
    ('Prayer requests table', 'Missing', 'Stored in localStorage only', 'N/A'),
    ('Hymns table', 'Missing', 'Hardcoded in JS, not in DB', 'hymns.js'),
    ('AI history table', 'Missing', 'Stored in localStorage only', 'N/A'),
    ('User settings table', 'Missing', 'Stored in localStorage only', 'N/A'),
    ('Subscriptions table', 'Missing', 'No billing system', 'N/A'),
    ('Feature flags table', 'Missing', 'No feature store', 'N/A'),
    ('Notifications table', 'Missing', 'No notification system', 'N/A'),
]
add_table(['Requirement', 'Status', 'Evidence', 'Location'], db_items)

doc.add_heading('3.4 Multi-Church Support', level=2)
church_items = [
    ('Church model', 'Missing', 'No database table, no API', 'N/A'),
    ('Ministry model', 'Missing', 'No database table, no API', 'N/A'),
    ('Publisher model', 'Missing', 'No publishing system', 'N/A'),
    ('Denominations', 'Missing', 'No denomination data', 'N/A'),
    ('Following churches', 'Missing', 'No follow/unfollow mechanism', 'N/A'),
    ('Church profiles', 'Missing', 'No church profile UI', 'N/A'),
    ('Church branding', 'Missing', 'No branding system', 'N/A'),
    ('Church-admin dashboard', 'Missing', 'No admin panel', 'N/A'),
    ('Multi-tenant architecture', 'Missing', 'Single-tenant only', 'N/A'),
    ('Devotional publishing', 'Partial', '12 churches hardcoded in JS; no DB, no API, no admin', 'churchDevotionals.js'),
]
add_table(['Requirement', 'Status', 'Evidence', 'Location'], church_items)

doc.add_heading('3.5 Offline-first Architecture', level=2)
offline_items = [
    ('Local database', 'Partial', 'localStorage used; no IndexedDB, no SQLite', 'App.jsx:304-335'),
    ('Caching', 'Partial', 'Bible chapters cached; no image cache', 'vite.config.js:37-45'),
    ('Offline queue', 'Missing', 'No queue for pending operations', 'N/A'),
    ('Delta sync', 'Missing', 'Full data push/pull only', 'sync.py:27-64'),
    ('Conflict resolution', 'Partial', 'Last-write-wins; no conflict UI', 'sync.py:47-64'),
    ('Sync status indicators', 'Missing', 'No sync status UI', 'N/A'),
]
add_table(['Requirement', 'Status', 'Evidence', 'Location'], offline_items)

doc.add_heading('3.6 Bible Module', level=2)
bible_items = [
    ('12 translations', 'Full', 'KJV + 11 via LLM', 'App.jsx:50-63, index.py:201-267'),
    ('Bible search', 'Partial', 'Concordance via AI; no text search', 'index.py:350-370'),
    ('Notes per chapter', 'Missing', 'No notes feature', 'N/A'),
    ('Bookmarks', 'Missing', 'No bookmark feature', 'N/A'),
    ('Cross references', 'Missing', 'No cross-reference data', 'N/A'),
    ('Verse sharing', 'Missing', 'No share functionality', 'N/A'),
    ('Highlighting', 'Missing', 'No highlight feature', 'N/A'),
    ('Reading history', 'Full', 'recentReads tracked', 'App.jsx:246, syncService.js:47'),
]
add_table(['Requirement', 'Status', 'Evidence', 'Location'], bible_items)

doc.add_heading('3.7 AI Assistant', level=2)
ai_items = [
    ('RAG implementation', 'Full', 'Pinecone vector search + LLM generation', 'rag.py:91-148'),
    ('Prompt engineering', 'Partial', 'Basic system prompts; no verse grounding', 'index.py:279-300'),
    ('Source citations', 'Missing', 'No citations in AI responses', 'N/A'),
    ('Verse grounding', 'Missing', 'AI not grounded in scripture', 'N/A'),
    ('Confidence scores', 'Missing', 'No confidence indicators', 'N/A'),
    ('Conversation history', 'Partial', 'localStorage only; no server-side', 'App.jsx:231'),
    ('Multi-provider support', 'Full', 'GROQ, OpenAI, OpenRouter', 'llm_provider.py:7-54'),
    ('Rate limiting', 'Missing', 'No rate limiting on any endpoint', 'N/A'),
]
add_table(['Requirement', 'Status', 'Evidence', 'Location'], ai_items)

doc.add_heading('3.8 Prayer Module', level=2)
prayer_items = [
    ('Prayer requests', 'Partial', 'Log minutes only; no request text', 'App.jsx:606-614'),
    ('Categories', 'Missing', 'No prayer categories', 'N/A'),
    ('Prayer streaks', 'Full', 'getStreak() calculates consecutive days', 'App.jsx:176-186'),
    ('Answer tracking', 'Missing', 'No answer/praise reporting', 'N/A'),
    ('Statistics', 'Missing', 'No prayer analytics', 'N/A'),
]
add_table(['Requirement', 'Status', 'Evidence', 'Location'], prayer_items)

doc.add_heading('3.9 Journal', level=2)
journal_items = [
    ('CRUD operations', 'Full', 'Create, read, edit, delete diary entries', 'App.jsx:631-662'),
    ('Encryption', 'Missing', 'Plaintext in localStorage', 'N/A'),
    ('Markdown support', 'Missing', 'Plain text only', 'N/A'),
    ('Search', 'Missing', 'No search in diary', 'N/A'),
    ('Export', 'Partial', 'Full backup export; no diary-only export', 'App.jsx:920-955'),
    ('Local storage', 'Full', 'localStorage persistence', 'App.jsx:327'),
    ('Cloud sync', 'Full', 'Synced via user_data table', 'syncService.js:41'),
]
add_table(['Requirement', 'Status', 'Evidence', 'Location'], journal_items)

doc.add_heading('3.10 Hymn Module', level=2)
hymn_items = [
    ('1,001 hymns', 'Full', '1,001 hymns with title, author, lyrics, category', 'hymns.js:1-1003'),
    ('Audio playback', 'Partial', '54 tunes via Web Audio API; fallback for ~100', 'hymnMusic.js:1-123'),
    ('Favorites', 'Full', 'Toggle favorite, persist to localStorage', 'App.jsx:842-848'),
    ('Search', 'Full', 'Search by title, author, category', 'App.jsx:860-870'),
    ('Offline playback', 'Missing', 'Tunes fetched from API; fallback only', 'hymnMusic.js:62-72'),
]
add_table(['Requirement', 'Status', 'Evidence', 'Location'], hymn_items)

doc.add_heading('3.11 Personalization', level=2)
personal_items = [
    ('Rearrange dashboard', 'Full', 'Drag-and-drop nav reorder', 'App.jsx:710-775'),
    ('Hide features', 'Missing', 'No hide/show feature toggle', 'N/A'),
    ('Select denominations', 'Missing', 'No denomination selection', 'N/A'),
    ('Choose Bible versions', 'Full', '12 translations selectable', 'App.jsx:50-63'),
    ('Choose devotionals', 'Full', '12 churches selectable', 'App.jsx:1675-1795'),
    ('Select themes', 'Full', '5 themes + custom colors', 'App.jsx:145-169'),
    ('Configure notifications', 'Partial', 'Toggles exist; no actual implementation', 'App.jsx:1954-1970'),
]
add_table(['Requirement', 'Status', 'Evidence', 'Location'], personal_items)

doc.add_heading('3.12 Feature Store', level=2)
feature_items = [
    ('Feature flags', 'Missing', 'No feature flag system', 'N/A'),
    ('Remote configuration', 'Missing', 'No remote config', 'N/A'),
    ('Premium gating', 'Partial', 'Auth-based premium; no subscription check', 'App.jsx:203'),
    ('Dynamic modules', 'Missing', 'No dynamic module loading', 'N/A'),
]
add_table(['Requirement', 'Status', 'Evidence', 'Location'], feature_items)

doc.add_heading('3.13 Feedback System', level=2)
feedback_items = [
    ('Bug reports', 'Missing', 'No in-app bug reporting', 'N/A'),
    ('Feature requests', 'Missing', 'No feature request mechanism', 'N/A'),
    ('Ratings', 'Missing', 'No rating system', 'N/A'),
    ('NPS', 'Missing', 'No NPS surveys', 'N/A'),
    ('Admin review', 'Missing', 'No admin dashboard', 'N/A'),
]
add_table(['Requirement', 'Status', 'Evidence', 'Location'], feedback_items)

doc.add_heading('3.14 Subscription System', level=2)
sub_items = [
    ('Billing integration', 'Missing', 'No payment provider configured', 'N/A'),
    ('Premium plans', 'Partial', 'Price shown ($4.99/mo) but no actual billing', 'PremiumGate.jsx:46'),
    ('Free plans', 'Full', 'All features free for logged-in users', 'App.jsx:203'),
    ('Purchase validation', 'Missing', 'No purchase flow', 'N/A'),
    ('Restore purchases', 'Missing', 'No restore mechanism', 'N/A'),
    ('Feature gating', 'Partial', 'Auth-based; not subscription-based', 'App.jsx:203'),
]
add_table(['Requirement', 'Status', 'Evidence', 'Location'], sub_items)

doc.add_heading('3.15 Community', level=2)
community_items = [
    ('Forum', 'Missing', 'No forum feature', 'N/A'),
    ('Posts & replies', 'Missing', 'No posting system', 'N/A'),
    ('Moderation', 'Missing', 'No moderation tools', 'N/A'),
    ('Reports', 'Missing', 'No reporting system', 'N/A'),
    ('Likes', 'Missing', 'No engagement features', 'N/A'),
]
add_table(['Requirement', 'Status', 'Evidence', 'Location'], community_items)

doc.add_heading('3.16 Notifications', level=2)
notif_items = [
    ('Push notifications', 'Missing', 'No FCM/APNs integration', 'N/A'),
    ('Local notifications', 'Missing', 'No Notification API usage', 'N/A'),
    ('Scheduled reminders', 'Missing', 'No scheduling', 'N/A'),
    ('Devotional reminders', 'Missing', 'No daily reminders', 'N/A'),
    ('Prayer reminders', 'Missing', 'No prayer reminders', 'N/A'),
]
add_table(['Requirement', 'Status', 'Evidence', 'Location'], notif_items)

doc.add_page_break()

# ============================================================
# 4. MISSING FEATURES REPORT
# ============================================================
doc.add_heading('4. Missing Features Report', level=1)

doc.add_heading('4.1 Critical Missing Features', level=2)
doc.add_paragraph('These features block production readiness:')

critical_missing = [
    ('Password Reset', 'No mechanism to recover lost passwords', 'auth.py', 'Week 1-2'),
    ('Email Verification', 'No email validation on signup', 'auth.py', 'Week 1-2'),
    ('Account Deletion', 'No endpoint or UI to delete account', 'auth.py, App.jsx', 'Week 1'),
    ('Rate Limiting', 'No rate limiting on any endpoint', 'index.py', 'Week 1'),
    ('Input Sanitization', 'No XSS protection beyond React escaping', 'index.py', 'Week 1'),
    ('CORS Restriction', 'CORS allows all origins with credentials', 'index.py:38', 'Week 1'),
    ('Secrets Management', 'Hardcoded API keys in reingest.py', 'reingest.py:4,7', 'Immediate'),
]
add_table(['Feature', 'Impact', 'File', 'Effort'], critical_missing)

doc.add_heading('4.2 High Priority Missing Features', level=2)
high_missing = [
    ('Subscription Billing', 'No revenue generation', 'New: billing module', 'Week 3-4'),
    ('Push Notifications', 'No user re-engagement', 'New: notification service', 'Week 3-4'),
    ('Church Model / Directory', 'PRD requirement', 'New: church module', 'Week 4-6'),
    ('Feature Flags', 'No dynamic feature control', 'New: feature store', 'Week 2-3'),
    ('Privacy Dashboard', 'GDPR compliance', 'New: privacy UI', 'Week 2'),
    ('Consent Management', 'Legal compliance', 'New: consent flow', 'Week 2'),
    ('Backend Tests', 'No API test coverage', 'New: pytest tests', 'Week 2-3'),
    ('E2E Tests', 'No integration test coverage', 'New: Playwright tests', 'Week 3-4'),
    ('CI/CD Pipeline', 'No automated testing on push', '.github/workflows/', 'Week 1'),
]
add_table(['Feature', 'Impact', 'File', 'Effort'], high_missing)

doc.add_heading('4.3 Medium Priority Missing Features', level=2)
medium_missing = [
    ('Offline Queue', 'No pending operations queue', 'New: queue service', 'Week 4-5'),
    ('Conflict Resolution UI', 'No user-facing conflict handling', 'New: conflict UI', 'Week 5'),
    ('Bible Notes', 'No per-chapter notes', 'New: notes feature', 'Week 5-6'),
    ('Bible Bookmarks', 'No verse bookmarking', 'New: bookmarks feature', 'Week 5-6'),
    ('Verse Sharing', 'No share functionality', 'New: share feature', 'Week 6'),
    ('Prayer Categories', 'No categorization', 'New: categories', 'Week 6'),
    ('Prayer Answer Tracking', 'No praise reports', 'New: answer tracking', 'Week 6'),
    ('Diary Encryption', 'Plaintext diary entries', 'New: encryption layer', 'Week 7'),
    ('Diary Markdown', 'No rich text support', 'New: markdown editor', 'Week 7'),
    ('Diary Search', 'No search in journal', 'New: search feature', 'Week 7'),
    ('Admin Dashboard', 'No content moderation', 'New: admin panel', 'Week 8'),
]
add_table(['Feature', 'Impact', 'File', 'Effort'], medium_missing)

doc.add_page_break()

# ============================================================
# 5. SECURITY AUDIT
# ============================================================
doc.add_heading('5. Security Audit', level=1)

doc.add_heading('5.1 Critical Security Issues', level=2)
sec_critical = [
    ('SEC-001', 'Critical', 'Hardcoded API keys in reingest.py', 'reingest.py:4,7', 'Rotate keys immediately; use env vars'),
    ('SEC-002', 'Critical', 'Hardcoded JWT secret fallback', 'auth.py:19', 'Remove fallback; require env var'),
    ('SEC-003', 'Critical', 'No auth on /api/rag/search, /api/rag/ingest, /api/llm/chat', 'index.py:179,184,189', 'Add auth dependency'),
    ('SEC-004', 'High', 'CORS allow_origins=["*"] with credentials', 'index.py:38', 'Restrict to specific origins'),
    ('SEC-005', 'High', 'get_connection() bypasses pool', 'database.py:30', 'Migrate to pool.acquire()'),
    ('SEC-006', 'High', '30-day JWT token expiry', 'auth.py:21', 'Reduce to 24h + refresh tokens'),
    ('SEC-007', 'High', 'Capacitor keystore password in config', 'capacitor.config.json:14', 'Move to env var or secure storage'),
    ('SEC-008', 'Medium', 'No rate limiting on any endpoint', 'index.py', 'Add rate limiting middleware'),
    ('SEC-009', 'Medium', 'No input validation on LLM messages', 'llm_provider.py:13', 'Validate message content'),
    ('SEC-010', 'Medium', 'No RBAC - all users equal', 'auth.py', 'Add role system'),
    ('SEC-011', 'Low', 'No CSRF protection', 'index.py', 'Add CSRF tokens for state-changing ops'),
    ('SEC-012', 'Low', 'No Content-Security-Policy headers', 'index.py', 'Add CSP headers'),
]
add_table(['ID', 'Severity', 'Issue', 'Location', 'Remediation'], sec_critical)

doc.add_heading('5.2 Authentication Security', level=2)
auth_sec = [
    ('Password hashing', 'bcrypt via passlib', 'Strong', 'auth.py:55'),
    ('JWT algorithm', 'HS256', 'Acceptable', 'auth.py:20'),
    ('Token expiry', '30 days', 'Too long', 'auth.py:21'),
    ('Password min length', '6 characters', 'Weak (recommend 8+)', 'auth.py:32'),
    ('Brute force protection', 'None', 'Missing', 'N/A'),
    ('Account lockout', 'None', 'Missing', 'N/A'),
    ('Session invalidation', 'Token-based; no revocation', 'Weak', 'N/A'),
]
add_table(['Aspect', 'Implementation', 'Assessment', 'Location'], auth_sec)

doc.add_heading('5.3 Secrets in Source Code', level=2)
doc.add_paragraph(
    'The following secrets were found committed to source code and MUST be rotated immediately:'
)
secrets = [
    ('reingest.py:4', 'Pinecone API key', '[REDACTED — was in source, now uses env var]'),
    ('reingest.py:7', 'OpenAI API key', '[REDACTED — was in source, now uses env var]'),
    ('auth.py:19', 'JWT secret', '[REDACTED — was fallback, now requires env var]'),
    ('capacitor.config.json:14', 'Keystore password', '[REDACTED]'),
    ('.env:6-7', 'API keys (commented)', '[REDACTED — .env not committed]'),
]
add_table(['Location', 'Secret Type', 'Value (truncated)'], secrets)

doc.add_page_break()

# ============================================================
# 6. DATABASE AUDIT
# ============================================================
doc.add_heading('6. Database Audit', level=1)

doc.add_heading('6.1 Current Schema', level=2)
doc.add_paragraph('Only 2 tables exist in the database:')

schema = [
    ('users', 'id (PK), email (UNIQUE), name, password_hash, plan, created_at', 'database.py:37-44'),
    ('user_data', 'id (PK), user_id (FK), data_type, data (JSONB), updated_at; UNIQUE(user_id, data_type)', 'database.py:53-62'),
]
add_table(['Table', 'Columns', 'Location'], schema)

doc.add_heading('6.2 Missing Tables (per PRD)', level=2)
missing_tables = [
    ('churches', 'Church profiles, branding, denomination', 'Multi-church support'),
    ('ministries', 'Ministry organization', 'Multi-church support'),
    ('subscriptions', 'User subscription plans, billing', 'Payment system'),
    ('notifications', 'Push notification tokens, preferences', 'Notification system'),
    ('feature_flags', 'Feature toggles, A/B tests', 'Feature store'),
    ('prayer_requests', 'Prayer request text, categories, answers', 'Prayer module'),
    ('bible_notes', 'Per-chapter user notes', 'Bible module'),
    ('bible_bookmarks', 'Saved verses', 'Bible module'),
    ('ai_conversations', 'Server-side chat history', 'AI assistant'),
    ('feedback', 'Bug reports, feature requests', 'Feedback system'),
    ('groups', 'Small group management', 'Community features'),
    ('forum_posts', 'Community forum content', 'Community features'),
]
add_table(['Table', 'Purpose', 'PRD Section'], missing_tables)

doc.add_heading('6.3 Schema Issues', level=2)
schema_issues = [
    ('Redundant ALTER TABLE', 'database.py:46-52 adds plan column that already exists in CREATE TABLE', 'Remove redundant migration'),
    ('No created_at on user_data', 'Only updated_at; no creation timestamp', 'Add created_at column'),
    ('No indexes on user_data.data_type', 'Queries filter by data_type but no index', 'Add composite index'),
    ('No soft delete', 'Hard delete only; no audit trail', 'Add deleted_at column'),
    ('JSONB data_type allows anything', 'No validation on JSON structure', 'Add CHECK constraints'),
]
add_table(['Issue', 'Description', 'Fix'], schema_issues)

doc.add_page_break()

# ============================================================
# 7. API AUDIT
# ============================================================
doc.add_heading('7. API Audit', level=1)

doc.add_heading('7.1 Endpoint Inventory (24 total)', level=2)
endpoints = [
    ('GET', '/api/health', 'No', 'Health check', 'index.py:86'),
    ('GET', '/api/dbtest', 'No', 'DB connection test', 'index.py:95'),
    ('GET', '/api/pinetest', 'No', 'Pinecone test', 'index.py:107'),
    ('POST', '/api/auth/register', 'No', 'User registration', 'index.py:120'),
    ('POST', '/api/auth/login', 'No', 'User login', 'index.py:132'),
    ('POST', '/api/auth/google', 'No', 'Google OAuth', 'index.py:144'),
    ('POST', '/api/auth/legal-accept', 'Yes', 'Accept legal terms', 'index.py:149'),
    ('GET', '/api/auth/legal-status', 'Yes', 'Legal status', 'index.py:159'),
    ('GET', '/api/sync/pull', 'Yes', 'Pull user data', 'index.py:169'),
    ('POST', '/api/sync/push', 'Yes', 'Push user data', 'index.py:174'),
    ('POST', '/api/rag/search', 'No', 'RAG vector search', 'index.py:179'),
    ('POST', '/api/rag/ingest', 'No', 'Ingest to Pinecone', 'index.py:184'),
    ('POST', '/api/llm/chat', 'No', 'Raw LLM chat', 'index.py:189'),
    ('GET', '/api/llm/providers', 'No', 'List providers', 'index.py:196'),
    ('GET', '/api/bible/versions', 'No', 'Bible versions', 'index.py:201'),
    ('GET', '/api/bible', 'No', 'Fetch chapter', 'index.py:267'),
    ('POST', '/api/chat', 'No', 'AI chatbot', 'index.py:279'),
    ('POST', '/api/bible/explain', 'No', 'Explain verse', 'index.py:301'),
    ('POST', '/api/bible/commentary', 'No', 'Commentary', 'index.py:324'),
    ('POST', '/api/bible/concordance', 'No', 'Concordance search', 'index.py:350'),
    ('POST', '/api/hymns/explain', 'No', 'Hymn explanation', 'index.py:371'),
    ('POST', '/api/devotional/generate', 'No', 'Generate devotional', 'index.py:394'),
    ('GET', '/api/hymns/tune/{id}', 'No', 'Hymn melody', 'index.py:415'),
    ('POST', '/api/bible/compare', 'No', 'Compare translations', 'index.py:427'),
]
add_table(['Method', 'Path', 'Auth', 'Description', 'Location'], endpoints)

doc.add_heading('7.2 API Issues', level=2)
api_issues = [
    ('No versioning', 'All endpoints at /api/*; no /api/v1/*', 'Add URL-based versioning', 'High'),
    ('No rate limiting', 'All endpoints unprotected', 'Add rate limit middleware', 'Critical'),
    ('No request validation', 'LLM messages untyped', 'Add Pydantic validation', 'Medium'),
    ('No response caching', 'Bible versions fetched every time', 'Add Cache-Control headers', 'Low'),
    ('Duplicate endpoints', '/api/chat and /api/llm/chat both exist', 'Consolidate', 'Medium'),
    ('No health details', '/api/health returns minimal info', 'Add DB/Pinecone status', 'Low'),
    ('Silent failures', 'rag_search swallows exceptions', 'Return error responses', 'Medium'),
]
add_table(['Issue', 'Description', 'Fix', 'Priority'], api_issues)

doc.add_page_break()

# ============================================================
# 8. ARCHITECTURE REVIEW
# ============================================================
doc.add_heading('8. Architecture Review', level=1)

doc.add_heading('8.1 Current Architecture', level=2)
arch = [
    ('Frontend', 'React 19 SPA with localStorage', 'Single-file App.jsx (2,278 lines)'),
    ('Backend', 'Python FastAPI on Vercel Serverless', '4 modules: auth, sync, rag, llm_provider'),
    ('Database', 'PostgreSQL on Aiven', '2 tables, connection pool'),
    ('Vector Store', 'Pinecone', '54 Bible verses, 1024-dim'),
    ('AI', 'GROQ/OpenAI/OpenRouter', 'Multi-provider with fallback'),
    ('Auth', 'JWT + Google OAuth', '30-day tokens, bcrypt passwords'),
    ('Deployment', 'Vercel (frontend + backend)', 'GitHub Actions for backend only'),
]
add_table(['Layer', 'Technology', 'Notes'], arch)

doc.add_heading('8.2 Architecture Issues', level=2)
arch_issues = [
    ('Monolithic frontend', 'App.jsx is 2,278 lines with 73 state variables', 'Split into route-based components'),
    ('No state management', 'All state in App.jsx; no Context/Redux', 'Add React Context or Zustand'),
    ('No API abstraction', 'fetch calls scattered throughout', 'Create API service layer'),
    ('No error boundaries', 'No React error boundary component', 'Add ErrorBoundary wrapper'),
    ('No loading states', 'Inconsistent loading indicators', 'Standardize loading patterns'),
    ('No i18n', 'All strings hardcoded in English', 'Add react-i18next'),
    ('No analytics', 'No usage tracking', 'Add privacy-respecting analytics'),
    ('No monitoring', 'No error tracking (Sentry etc.)', 'Add error monitoring'),
]
add_table(['Issue', 'Description', 'Recommendation'], arch_issues)

doc.add_page_break()

# ============================================================
# 9. CODE QUALITY REPORT
# ============================================================
doc.add_heading('9. Code Quality Report', level=1)

doc.add_heading('9.1 Frontend Code Quality', level=2)
fe_quality = [
    ('Component size', 'App.jsx: 2,278 lines, 73 state vars', 'Critical', 'Split into 10+ components'),
    ('Single responsibility', 'One file handles everything', 'Critical', 'Extract feature modules'),
    ('Dead code', 'PremiumGate imported but unused', 'Medium', 'Remove or integrate'),
    ('Dead code', 'selectedTimezone state unused', 'Low', 'Remove'),
    ('Code duplication', 'Multiple相似的 useEffect patterns', 'Medium', 'Extract custom hooks'),
    ('Naming conventions', 'Consistent camelCase/PascalCase', 'Good', 'N/A'),
    ('CSS organization', 'Single App.css (1,950+ lines)', 'Medium', 'Split into module CSS'),
    ('Error handling', 'Try/catch with generic messages', 'Medium', 'Add specific error types'),
]
add_table(['Aspect', 'Finding', 'Severity', 'Recommendation'], fe_quality)

doc.add_heading('9.2 Backend Code Quality', level=2)
be_quality = [
    ('Module structure', '4 modules; clean separation', 'Good', 'N/A'),
    ('Code duplication', 'call_llm() and call_llm_multi() nearly identical', 'Medium', 'Refactor to shared helper'),
    ('Dead code', 'optional_user() never called', 'Low', 'Remove or use'),
    ('Dead code', 'get_embedding() in llm_provider.py never called', 'Low', 'Remove'),
    ('Dead code', 'PINECONE_HOST defined but unused', 'Low', 'Remove'),
    ('Error handling', 'Bare except Exception in google_auth', 'Medium', 'Catch specific exceptions'),
    ('Type hints', 'Partial; some functions lack return types', 'Low', 'Add complete type hints'),
    ('Docstrings', 'No docstrings on any function', 'Low', 'Add docstrings'),
]
add_table(['Aspect', 'Finding', 'Severity', 'Recommendation'], be_quality)

doc.add_heading('9.3 Test Quality', level=2)
test_quality = [
    ('Unit test count', '61 tests passing', 'Low coverage', 'Need 200+ tests'),
    ('Test files', '2 files (appUtils, dateUtils)', 'Minimal', 'Add tests for all modules'),
    ('Backend tests', '0 tests', 'Critical', 'Add pytest tests for all endpoints'),
    ('Integration tests', '0 tests', 'Critical', 'Add API integration tests'),
    ('E2E tests', '0 tests', 'High', 'Add Playwright tests'),
    ('Coverage reporting', 'Not configured', 'Medium', 'Add coverage reporting'),
    ('Test automation', 'No CI integration', 'High', 'Add to GitHub Actions'),
]
add_table(['Aspect', 'Finding', 'Severity', 'Recommendation'], test_quality)

doc.add_page_break()

# ============================================================
# 10. PERFORMANCE REVIEW
# ============================================================
doc.add_heading('10. Performance Review', level=1)

doc.add_heading('10.1 Bundle Analysis', level=2)
bundle = [
    ('Total bundle size', '~1,057 KB (minified)', 'Above 500KB recommendation'),
    ('CSS', '~90 KB', 'Acceptable'),
    ('JavaScript', '~1,057 KB', 'Needs code splitting'),
    ('Hymns data', '~421 KB embedded', 'Should lazy-load'),
    ('Devotionals data', '~315 KB embedded', 'Should lazy-load'),
    ('PWA precache', '~1,592 KB', 'Large for mobile'),
]
add_table(['Metric', 'Value', 'Assessment'], bundle)

doc.add_heading('10.2 Performance Issues', level=2)
perf_issues = [
    ('No code splitting', 'All code in single bundle', 'Use React.lazy() for views'),
    ('No lazy loading', 'All data loaded on startup', 'Load hymns/devotionals on demand'),
    ('No image optimization', 'PNG icons only', 'Use WebP with sharp'),
    ('No virtual scrolling', '1,001 hymns rendered at once', 'Use windowed list'),
    ('Large localStorage', '14 keys with potential 5MB limit', 'Implement cleanup/ archival'),
    ('Synchronous embedding', 'rag.py blocks event loop', 'Make async'),
    ('No API caching', 'Bible versions fetched every request', 'Add Cache-Control'),
    ('No CDN', 'All assets from Vercel', 'Add Cloudflare for static assets'),
]
add_table(['Issue', 'Description', 'Recommendation'], perf_issues)

doc.add_page_break()

# ============================================================
# 11. TECHNICAL DEBT REPORT
# ============================================================
doc.add_heading('11. Technical Debt Report', level=1)

doc.add_heading('11.1 Critical Debt', level=2)
crit_debt = [
    ('Hardcoded secrets', 'API keys in source code', 'Immediate', 'Rotate + use env vars'),
    ('Monolithic App.jsx', '2,278 lines, 73 state vars', '1-2 weeks', 'Split into components'),
    ('No backend tests', 'Zero test coverage for API', '1 week', 'Add pytest suite'),
    ('No CI/CD for frontend', 'Only backend deploy workflow', '2-3 days', 'Add lint+test+build workflow'),
    ('Default README', 'Vite template README', '1 hour', 'Write project README'),
]
add_table(['Debt', 'Impact', 'Effort to Fix', 'Recommendation'], crit_debt)

doc.add_heading('11.2 High Priority Debt', level=2)
high_debt = [
    ('Duplicate LLM functions', 'call_llm/call_llm_multi nearly identical', '2 hours', 'Refactor to shared helper'),
    ('Dead code in backend', 'optional_user, get_embedding unused', '1 hour', 'Remove'),
    ('Synchronous embedding', 'Blocks event loop', '4 hours', 'Make async'),
    ('Silent error swallowing', 'rag_search, fetch_bible_ai hide errors', '4 hours', 'Add error propagation'),
    ('No error boundaries', 'React errors crash entire app', '2 hours', 'Add ErrorBoundary'),
    ('Inconsistent connection mgmt', 'get_connection vs pool.acquire', '4 hours', 'Standardize on pool'),
]
add_table(['Debt', 'Impact', 'Effort to Fix', 'Recommendation'], high_debt)

doc.add_heading('11.3 Medium Priority Debt', level=2)
med_debt = [
    ('No i18n', 'All strings in English', '2-3 weeks', 'Add react-i18next'),
    ('No state management', 'Props drilling through App.jsx', '1-2 weeks', 'Add React Context'),
    ('Single CSS file', '1,950+ lines', '1 week', 'Split into CSS modules'),
    ('No API versioning', 'All endpoints unversioned', '1 day', 'Add /api/v1/ prefix'),
    ('No CSP headers', 'No Content-Security-Policy', '2 hours', 'Add security headers'),
    ('Unused imports', 'Request, JSONResponse in index.py', '15 min', 'Remove'),
]
add_table(['Debt', 'Impact', 'Effort to Fix', 'Recommendation'], med_debt)

doc.add_page_break()

# ============================================================
# 12. REMAINING IMPLEMENTATION CHECKLIST
# ============================================================
doc.add_heading('12. Remaining Implementation Checklist', level=1)

doc.add_heading('12.1 Critical (Must Complete Before Production)', level=2)
checklist_crit = [
    ('Rotate all hardcoded API keys', 'reingest.py, .env', 'Immediate'),
    ('Remove JWT secret fallback', 'auth.py:19', 'Immediate'),
    ('Add auth to RAG/LLM endpoints', 'index.py:179,184,189', 'Day 1'),
    ('Restrict CORS origins', 'index.py:38', 'Day 1'),
    ('Fix connection pool bypass', 'database.py:30', 'Day 1'),
    ('Add rate limiting', 'index.py', 'Week 1'),
    ('Add password reset flow', 'auth.py + frontend', 'Week 1-2'),
    ('Add email verification', 'auth.py + frontend', 'Week 1-2'),
    ('Add account deletion endpoint', 'auth.py', 'Week 1'),
    ('Add backend test suite', 'New: tests/', 'Week 1-2'),
    ('Add CI/CD for frontend', '.github/workflows/', 'Week 1'),
    ('Write project README', 'README.md', 'Day 1'),
]
add_table(['Task', 'Location', 'Effort'], checklist_crit)

doc.add_heading('12.2 High Priority (Should Complete)', level=2)
checklist_high = [
    ('Implement subscription billing', 'New: billing module', 'Week 3-4'),
    ('Add push notifications', 'New: notification service', 'Week 3-4'),
    ('Add church model + directory', 'New: church module', 'Week 4-6'),
    ('Add feature flags', 'New: feature store', 'Week 2-3'),
    ('Add privacy dashboard', 'New: privacy UI', 'Week 2'),
    ('Add consent management', 'New: consent flow', 'Week 2'),
    ('Add E2E tests', 'New: Playwright tests', 'Week 3-4'),
    ('Split App.jsx into components', 'src/components/', 'Week 2-3'),
    ('Add React Context for state', 'New: contexts/', 'Week 2-3'),
    ('Add error boundaries', 'New: ErrorBoundary.jsx', 'Week 1'),
]
add_table(['Task', 'Location', 'Effort'], checklist_high)

doc.add_page_break()

# ============================================================
# 13. PRIORITIZED SPRINT PLAN
# ============================================================
doc.add_heading('13. Prioritized Sprint Plan', level=1)

doc.add_heading('Sprint 1: Security & Infrastructure (Week 1)', level=2)
sprint1 = [
    ('Day 1', 'Rotate all hardcoded API keys', 'Critical', '2h'),
    ('Day 1', 'Remove JWT secret fallback', 'Critical', '1h'),
    ('Day 1', 'Add auth to RAG/LLM endpoints', 'Critical', '2h'),
    ('Day 1', 'Restrict CORS origins', 'Critical', '1h'),
    ('Day 1', 'Fix connection pool bypass', 'Critical', '2h'),
    ('Day 1', 'Write project README', 'High', '2h'),
    ('Day 2', 'Add rate limiting middleware', 'Critical', '4h'),
    ('Day 2', 'Add error boundaries', 'High', '2h'),
    ('Day 3', 'Add backend test suite (10+ tests)', 'Critical', '6h'),
    ('Day 4', 'Add CI/CD workflow for frontend', 'High', '4h'),
    ('Day 5', 'Add password reset endpoint + UI', 'Critical', '6h'),
]
add_table(['Day', 'Task', 'Priority', 'Effort'], sprint1)

doc.add_heading('Sprint 2: Auth & Compliance (Week 2)', level=2)
sprint2 = [
    ('Day 1-2', 'Add email verification flow', 'Critical', '8h'),
    ('Day 2-3', 'Add account deletion endpoint + UI', 'Critical', '6h'),
    ('Day 3', 'Add privacy dashboard UI', 'High', '4h'),
    ('Day 3-4', 'Add consent management flow', 'High', '4h'),
    ('Day 4', 'Add CSP security headers', 'Medium', '2h'),
    ('Day 4-5', 'Refactor App.jsx into components', 'High', '8h'),
    ('Day 5', 'Add React Context for state', 'High', '4h'),
]
add_table(['Day', 'Task', 'Priority', 'Effort'], sprint2)

doc.add_heading('Sprint 3: Features & Testing (Week 3)', level=2)
sprint3 = [
    ('Day 1-2', 'Add feature flags system', 'High', '6h'),
    ('Day 2-3', 'Add push notification service', 'High', '8h'),
    ('Day 3-4', 'Add E2E tests (Playwright)', 'High', '8h'),
    ('Day 4-5', 'Add Bible notes feature', 'Medium', '6h'),
    ('Day 5', 'Add Bible bookmarks feature', 'Medium', '4h'),
]
add_table(['Day', 'Task', 'Priority', 'Effort'], sprint3)

doc.add_heading('Sprint 4: Payments & Community (Week 4)', level=2)
sprint4 = [
    ('Day 1-3', 'Integrate Flutterwave payments', 'High', '12h'),
    ('Day 3-4', 'Add church model + directory', 'High', '8h'),
    ('Day 4-5', 'Add prayer categories + answer tracking', 'Medium', '6h'),
    ('Day 5', 'Add verse sharing feature', 'Medium', '4h'),
]
add_table(['Day', 'Task', 'Priority', 'Effort'], sprint4)

doc.add_page_break()

# ============================================================
# 14. GITHUB ISSUES / TASKS
# ============================================================
doc.add_heading('14. Prioritized GitHub Issues / Tasks', level=1)

doc.add_heading('Critical Issues', level=2)
issues_crit = [
    ('SEC-001', 'Rotate hardcoded API keys in reingest.py', 'security', 'P0'),
    ('SEC-002', 'Remove JWT secret fallback in auth.py', 'security', 'P0'),
    ('SEC-003', 'Add authentication to RAG/LLM endpoints', 'security', 'P0'),
    ('SEC-004', 'Restrict CORS to specific origins', 'security', 'P0'),
    ('SEC-005', 'Fix connection pool bypass in database.py', 'security', 'P0'),
    ('AUTH-001', 'Implement password reset flow', 'enhancement', 'P0'),
    ('AUTH-002', 'Implement email verification', 'enhancement', 'P0'),
    ('AUTH-003', 'Implement account deletion endpoint', 'enhancement', 'P0'),
    ('INFRA-001', 'Add rate limiting to all endpoints', 'enhancement', 'P0'),
    ('TEST-001', 'Add backend test suite (pytest)', 'testing', 'P0'),
    ('CI-001', 'Add CI/CD workflow for frontend (lint+test+build)', 'infrastructure', 'P0'),
    ('DOC-001', 'Replace default Vite README with project README', 'documentation', 'P0'),
]
add_table(['ID', 'Title', 'Label', 'Priority'], issues_crit)

doc.add_heading('High Priority Issues', level=2)
issues_high = [
    ('FEAT-001', 'Implement subscription billing (Flutterwave)', 'enhancement', 'P1'),
    ('FEAT-002', 'Implement push notifications (FCM)', 'enhancement', 'P1'),
    ('FEAT-003', 'Add church model and directory', 'enhancement', 'P1'),
    ('FEAT-004', 'Add feature flags system', 'enhancement', 'P1'),
    ('FEAT-005', 'Add privacy dashboard UI', 'enhancement', 'P1'),
    ('FEAT-006', 'Add consent management flow', 'enhancement', 'P1'),
    ('TEST-002', 'Add E2E tests with Playwright', 'testing', 'P1'),
    ('REFACTOR-001', 'Split App.jsx into route-based components', 'refactoring', 'P1'),
    ('REFACTOR-002', 'Add React Context for state management', 'refactoring', 'P1'),
    ('SEC-006', 'Reduce JWT expiry to 24h + add refresh tokens', 'security', 'P1'),
]
add_table(['ID', 'Title', 'Label', 'Priority'], issues_high)

doc.add_heading('Medium Priority Issues', level=2)
issues_med = [
    ('FEAT-007', 'Add Bible notes per chapter', 'enhancement', 'P2'),
    ('FEAT-008', 'Add Bible bookmarks', 'enhancement', 'P2'),
    ('FEAT-009', 'Add verse sharing', 'enhancement', 'P2'),
    ('FEAT-010', 'Add prayer categories', 'enhancement', 'P2'),
    ('FEAT-011', 'Add prayer answer tracking', 'enhancement', 'P2'),
    ('FEAT-012', 'Add diary encryption', 'enhancement', 'P2'),
    ('FEAT-013', 'Add diary markdown support', 'enhancement', 'P2'),
    ('FEAT-014', 'Add diary search', 'enhancement', 'P2'),
    ('FEAT-015', 'Add offline queue for pending ops', 'enhancement', 'P2'),
    ('FEAT-016', 'Add conflict resolution UI', 'enhancement', 'P2'),
    ('REFACTOR-003', 'Add API versioning (/api/v1/)', 'refactoring', 'P2'),
    ('REFACTOR-004', 'Consolidate /api/chat and /api/llm/chat', 'refactoring', 'P2'),
    ('PERF-001', 'Add code splitting with React.lazy()', 'performance', 'P2'),
    ('PERF-002', 'Lazy-load hymns and devotionals data', 'performance', 'P2'),
    ('PERF-003', 'Add virtual scrolling for hymn list', 'performance', 'P2'),
]
add_table(['ID', 'Title', 'Label', 'Priority'], issues_med)

doc.add_page_break()

# ============================================================
# 15. FINAL RECOMMENDATION
# ============================================================
doc.add_heading('15. Final Recommendation', level=1)

doc.add_heading('Production Readiness Assessment', level=2)

p = doc.add_paragraph()
r = p.add_run('VERDICT: NOT PRODUCTION-READY')
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0xe7, 0x4c, 0x3c)

doc.add_paragraph()
doc.add_paragraph(
    'Based on a thorough audit of the entire BelieversFlow codebase against PRD v4.2.0, '
    'the application is assessed as BETA-READY but NOT PRODUCTION-READY.'
)

doc.add_heading('Justification', level=2)

reasons = [
    ('Security: 3 critical vulnerabilities', 'Hardcoded API keys, JWT secret fallback, and unauthenticated endpoints create immediate risk of data breach and cost abuse. These must be fixed before any public deployment.'),
    ('Authentication: Incomplete', 'No password reset, no email verification, no account deletion. Users cannot recover lost accounts or exercise GDPR rights.'),
    ('Testing: Minimal coverage', '61 unit tests cover only utility functions. Zero backend tests, zero integration tests, zero E2E tests. No confidence in API correctness.'),
    ('CI/CD: Incomplete', 'Only backend deployment workflow exists. No automated linting, testing, or building for frontend.'),
    ('Database: Schema incomplete', 'Only 2 tables. Missing churches, subscriptions, notifications, feature flags, and all community features.'),
    ('Infrastructure: No monitoring', 'No error tracking, no performance monitoring, no alerting.'),
    ('Compliance: Partial', 'Legal documents exist but no consent management UI, no privacy dashboard, no DPIA completed.'),
]

for title, desc in reasons:
    p = doc.add_paragraph()
    r = p.add_run(f'{title}: ')
    r.bold = True
    p.add_run(desc)

doc.add_heading('Recommended Path to Production', level=2)

path = [
    ('Week 1', 'Fix all critical security issues (rotate keys, add auth, restrict CORS)', 'MANDATORY'),
    ('Week 2', 'Complete auth flows (password reset, email verification, account deletion)', 'MANDATORY'),
    ('Week 2', 'Add backend test suite and CI/CD pipeline', 'MANDATORY'),
    ('Week 3-4', 'Implement subscription billing and push notifications', 'REQUIRED'),
    ('Week 4-6', 'Add church model, feature flags, privacy dashboard', 'REQUIRED'),
    ('Week 6-8', 'Complete missing features (notes, bookmarks, sharing, categories)', 'RECOMMENDED'),
    ('Week 8-10', 'Performance optimization, i18n, admin dashboard', 'RECOMMENDED'),
]
add_table(['Timeline', 'Action', 'Status'], path)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Estimated time to production readiness: 6-8 weeks of focused development.')
r.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run(
    'The application has a solid foundation with working core features (tasks, prayer, Bible, diary, '
    'hymns, devotionals, AI chat, cloud sync). The architectural decisions are sound. The primary '
    'gaps are in security hardening, testing infrastructure, and completing the PRD-mandated features. '
    'With the sprint plan outlined above, BelieversFlow can reach production readiness within the '
    'proposed timeline.'
)
r.italic = True

doc.add_paragraph()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run('— End of Audit Report —')
r.italic = True
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output = r'C:\Users\ibrah\Documents\Gemini\Christian_Todo\pitch\BelieversFlow_Audit_Report.docx'
doc.save(output)
print(f"Audit report saved to: {output}")
